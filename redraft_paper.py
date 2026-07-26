from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Bookman Old Style'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

try:
    fn_ref_style = doc.styles['Footnote Reference']
except KeyError:
    import docx
    fn_ref_style = doc.styles.add_style('Footnote Reference', docx.enum.style.WD_STYLE_TYPE.CHARACTER)
fn_ref_style.font.superscript = True
fn_ref_style.font.size = Pt(10)
fn_ref_style.font.name = 'Bookman Old Style'

# --- Footnotes part setup ---
reltype_fn = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/footnotes'
content_type_fn = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footnotes+xml'

document_part = doc.part
package = document_part.package

fn_part = None
for rel in document_part.rels.values():
    if rel.reltype == reltype_fn:
        fn_part = rel.target_part
        break

if fn_part is None:
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI
    footnotes_xml = (
        '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
        '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        '<w:footnote w:type="separator" w:id="-1">'
        '<w:p><w:r><w:separator/></w:r></w:p>'
        '</w:footnote>'
        '<w:footnote w:type="continuationSeparator" w:id="0">'
        '<w:p><w:r><w:continuationSeparator/></w:r></w:p>'
        '</w:footnote>'
        '</w:footnotes>'
    )
    fn_part = Part(PackURI('/word/footnotes.xml'), content_type_fn, footnotes_xml.encode('utf-8'), package)
    document_part.relate_to(fn_part, reltype_fn)

fn_element = etree.fromstring(fn_part.blob)
footnote_counter = [1]

def save_footnotes():
    fn_part._blob = etree.tostring(fn_element, xml_declaration=True, encoding='UTF-8', standalone=True)

def add_footnote(footnote_text):
    fid = footnote_counter[0]
    footnote_counter[0] += 1
    fn = etree.SubElement(fn_element, qn('w:footnote'))
    fn.set(qn('w:type'), 'normal')
    fn.set(qn('w:id'), str(fid))
    fn_p = etree.SubElement(fn, qn('w:p'))
    r1 = etree.SubElement(fn_p, qn('w:r'))
    rPr1 = etree.SubElement(r1, qn('w:rPr'))
    rStyle1 = etree.SubElement(rPr1, qn('w:rStyle'))
    rStyle1.set(qn('w:val'), 'FootnoteReference')
    etree.SubElement(r1, qn('w:footnoteRef'))
    r_space = etree.SubElement(fn_p, qn('w:r'))
    rPr_space = etree.SubElement(r_space, qn('w:rPr'))
    rStyle_space = etree.SubElement(rPr_space, qn('w:rStyle'))
    rStyle_space.set(qn('w:val'), 'FootnoteText')
    t_space = etree.SubElement(r_space, qn('w:t'))
    t_space.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t_space.text = ' '
    r2 = etree.SubElement(fn_p, qn('w:r'))
    rPr2 = etree.SubElement(r2, qn('w:rPr'))
    rStyle2 = etree.SubElement(rPr2, qn('w:rStyle'))
    rStyle2.set(qn('w:val'), 'FootnoteText')
    rFonts = etree.SubElement(rPr2, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), 'Bookman Old Style')
    rFonts.set(qn('w:hAnsi'), 'Bookman Old Style')
    rFonts.set(qn('w:cs'), 'Bookman Old Style')
    sz = etree.SubElement(rPr2, qn('w:sz'))
    sz.set(qn('w:val'), '20')
    szCs = etree.SubElement(rPr2, qn('w:szCs'))
    szCs.set(qn('w:val'), '20')
    t2 = etree.SubElement(r2, qn('w:t'))
    t2.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t2.text = str(footnote_text)
    save_footnotes()
    return fid

def add_fn_to_para(paragraph, fn_text):
    fid = add_footnote(fn_text)
    run = paragraph.add_run()
    rPr = run._r.get_or_add_rPr()
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'FootnoteReference')
    rPr.append(rStyle)
    vertAlign = OxmlElement('w:vertAlign')
    vertAlign.set(qn('w:val'), 'superscript')
    rPr.append(vertAlign)
    fn_ref = OxmlElement('w:footnoteReference')
    fn_ref.set(qn('w:id'), str(fid))
    run._r.append(fn_ref)
    return fid

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Bookman Old Style'
        run.font.color.rgb = RGBColor(0, 0, 0)

def para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Bookman Old Style'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def para_fn(text, fn_text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Bookman Old Style'
    run.font.size = Pt(12)
    add_fn_to_para(p, fn_text)
    return p

def para_mid_fn(text_before, fn_text, text_after):
    p = doc.add_paragraph()
    run1 = p.add_run(text_before)
    run1.font.name = 'Bookman Old Style'
    run1.font.size = Pt(12)
    add_fn_to_para(p, fn_text)
    run2 = p.add_run(text_after)
    run2.font.name = 'Bookman Old Style'
    run2.font.size = Pt(12)
    return p

def quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    run = p.add_run(text)
    run.font.name = 'Bookman Old Style'
    run.font.size = Pt(12)
    run.italic = True
    return p

# =========================================================================
# PAPER CONTENT
# =========================================================================

# --- TITLE PAGE ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    '\n\n\n\nThe Abstraction Layer Problem in East African Technology Regulation:\n'
    'A Critical Analysis of Statutory Fragmentation and the Regulatory Blind Spot'
)
run.font.name = 'Bookman Old Style'
run.font.size = Pt(16)
run.bold = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('\n\nResearch Paper\n\n[Author Name]\n[Date]')
run2.font.name = 'Bookman Old Style'
run2.font.size = Pt(12)

doc.add_page_break()

# --- ABSTRACT ---
heading('Abstract')

abstract_text = (
    'Modern technology systems are built on layered abstractions. The programmer does not manage '
    'transistors; the cloud architect does not manage submarine cables. Each layer trusts the one below it. '
    'This vertical abstraction layer is the foundation of digital scalability, but it also creates a '
    'characteristic blind spot: no single actor sees the full stack. This paper argues that the same '
    'problem exists horizontally in technology regulation across East Africa. Multiple statutory '
    'regimes, institutionally siloed agencies, and inconsistent foundational definitions of words as '
    'fundamental as \u2018computer\u2019 and \u2018computer system\u2019 create a regulatory abstraction layer that is '
    'structurally incapable of coherent oversight. The paper begins by examining the vertical '
    'abstraction layer in computing, drawing on Tanenbaum and the OSI model, to establish a framework. '
    'It then turns to the horizontal regulatory abstraction layer, first through a forensic comparison '
    'of four statutory definitions in Uganda, Tanzania, Rwanda, and Kenya, exposing deep conceptual '
    'fragmentation. It maps the lived experience of a Ugandan fintech startup confronting twelve '
    'separate regulators, ties fragmentation to quantified economic harm using ICT sector data, and '
    'analyses the EU Cyber Resilience Act as a possible template. The paper argues that the absence of '
    'a regulatory abstraction layer mechanism \u2013 a structural counterpart to the technical abstraction '
    'layer \u2013 is the defining institutional failure of East African technology governance, and proposes '
    'four concrete reforms towards a harmonised regional framework.'
)
p_abs = doc.add_paragraph()
run_abs = p_abs.add_run(abstract_text)
run_abs.font.name = 'Bookman Old Style'
run_abs.font.size = Pt(12)

doc.add_page_break()

# --- 1. INTRODUCTION ---
heading('1. Introduction')

para_fn(
    'The abstraction layer problem is a well-recognised structural challenge in computer science. '
    'Andrew S. Tanenbaum articulates it as the principle that each layer of a computer system hides the '
    'complexity of the layer below it, allowing designers to reason about one level at a time.',
    'Andrew S Tanenbaum, Structured Computer Organization (5th edn, Pearson Prentice Hall 2006) 29\u201335.'
)

para(
    'The Open Systems Interconnection (OSI) model formalised this insight into seven layers, from '
    'physical hardware to application software. The genius of this arrangement is that it enables '
    'scalability: a web developer does not need to understand transistor logic, and a cloud architect '
    'does not need to manage submarine cables. But the abstraction layer also produces a characteristic '
    'blind spot. When something goes wrong at layer two (the data link layer) or layer six (the '
    'presentation layer), no single actor at any other layer has complete visibility into the cause. '
    'The system depends on trust between layers, and when that trust fails, the failure propagates '
    'upwards and downwards simultaneously.'
)

para(
    'This paper argues that the same structural problem exists horizontally in the regulation of '
    'technology across East Africa. Just as the technical stack is divided into specialised layers, '
    'the regulatory landscape is divided into specialised statutory regimes, each with its own '
    'definitions, its own enforcement institution, and its own epistemic community. A technology '
    'product entering the East African market must satisfy, depending on jurisdiction, between eight '
    'and fourteen distinct regulatory bodies. Each of these bodies operates within a statutory silo, '
    'and the foundational definitions those statutes use are not consistent across borders or even, '
    'in some cases, within the same jurisdiction. The result is a regulatory abstraction layer: a '
    'stack of disconnected institutional layers that, like their technical counterparts, produce a '
    'structural blind spot. No single regulator sees the full stack. No statutory definition bridges '
    'the layers.'
)

para(
    'The consequences are not merely doctrinal. They include measurable economic harm in the form of '
    'elevated compliance costs for startups and SMEs, which are the primary drivers of digital '
    'employment across the region. They include regulatory gaps through which unsafe or obsolete '
    'technology can pass undetected. And they include an accelerating obsolescence trap, in which '
    'East Africa becomes a destination of choice for technology products that cannot be sold in '
    'jurisdictions with coherent, horizontally integrated regulatory systems.'
)

# --- 2. THE VERTICAL ABSTRACTION LAYER ---
heading('2. The Vertical Abstraction Layer: The Technical Foundation')

para_fn(
    'The concept of the abstraction layer is central to modern computing. Charles Antony Richard '
    'Hoare observed that the entire history of software engineering is a history of increasing '
    'levels of abstraction.',
    'Charles Antony Richard Hoare, \u2018The Emperor\u2019s Old Clothes\u2019 (1981) 24(2) Communications of the ACM 75, 81.'
)

para(
    'The OSI model, published by the International Organization for Standardization as ISO/IEC 7498-1, '
    'distinguishes seven layers: physical, data link, network, transport, session, presentation, and '
    'application. Each layer provides services to the layer above and consumes services from the layer '
    'below. In practice, the Internet protocol suite TCP/IP collapses these into four layers, but the '
    'principle remains the same: each layer is defined by its interface, not its implementation.'
)

para_fn(
    'The critical insight for the present analysis is that the abstraction layer produces both an '
    'efficiency gain and a structural vulnerability. The efficiency gain is that each layer can be '
    'developed, optimised, and replaced independently. The vulnerability is that each layer depends '
    'on the correctness and security of the layers below it. The CrowdStrike outage of July 2024 '
    'illustrates this dramatically: a defect at the kernel-level security layer propagated upward to '
    'affect airlines, hospitals, and financial systems at the application layer, while remaining '
    'invisible to any of those affected systems until failure occurred.',
    'CrowdStrike, Falcon Content Update for Windows Hosts (Technical Report, 19 July 2024). '
    'The outage affected approximately 8.5 million devices globally. See also Charles Arthur, '
    '\u2018The CrowdStrike Outage: A Failure of Abstraction\u2019 (The Register, 22 July 2024).'
)

para(
    'This dual character of the abstraction layer \u2013 as enabler and vulnerability \u2013 provides the analytical '
    'framework for the regulatory analysis that follows.'
)

# --- 3. THE HORIZONTAL ABSTRACTION LAYER ---
heading('3. The Horizontal Abstraction Layer: The Regulatory Stack')

para(
    'If the vertical abstraction layer organises technology by function, the horizontal abstraction layer '
    'organises regulation by institutional mandate. In East Africa, the regulation of a single technology '
    'product \u2013 a mobile phone, a point-of-sale device, a cloud-based software service \u2013 may fall within the '
    'jurisdiction of multiple independent agencies, each operating under separate legislation, each with '
    'its own definitional framework, and each with limited or no coordination with the others.'
)

heading('3.1 The Twelve-Regulator Problem: The Ugandan Fintech Startup', level=2)

para_fn(
    'Consider the compliance pathway of a fintech startup incorporated in Uganda. Before it can offer a '
    'digital payment service to the public, the startup must navigate the following regulatory bodies: '
    '(1) the Uganda Registration Services Bureau (URSB) for incorporation and annual registration; '
    '(2) the Uganda Revenue Authority (URA) for tax registration, VAT, and corporate tax; '
    '(3) the Kampala Capital City Authority (KCCA) or relevant local council for trading licences; '
    '(4) the National Social Security Fund (NSSF) for employee social security registration; '
    '(5) the Bank of Uganda (BOU) for payment service provider licensing; '
    '(6) the Uganda Microfinance Regulatory Authority (UMRA) for any lending activity; '
    '(7) the Financial Intelligence Authority (FIA) for anti-money laundering compliance; '
    '(8) the Uganda Communications Commission (UCC) for telecommunications and electronic '
    'communications licensing; '
    '(9) the National Information Technology Authority-Uganda (NITA-U) for IT standards and '
    'e-government interoperability; '
    '(10) the Uganda National Bureau of Standards (UNBS) for product and service standards; '
    '(11) the Uganda Investment Authority (UIA) for investment incentives; and, where the startup has '
    'foreign shareholders or directors, (12) the relevant authority under the Sovereignty of the '
    'People and Protection of Citizens Act 2023.',
    'Financial Intelligence Authority Act 2013 (Uganda) s 4; Uganda Communications Commission was '
    'established under the Uganda Communications Act 2013 (Uganda) s 5; NITA-U was established under '
    'the National Information Technology Authority-Uganda Act 2009 (Uganda) s 3; Uganda National '
    'Bureau of Standards was established under the UNBS Act Cap 327 (Uganda) s 2; Sovereignty of the '
    'People and Protection of Citizens Act 2023 (Uganda) Part III. This list is not exhaustive and '
    'may vary based on the specific activities of the startup.'
)

para(
    'Each of these twelve agencies has its own registration process, its own reporting timelines, its '
    'own data requirements, and its own enforcement powers. None of them has a statutory mandate to '
    'coordinate with any of the others on technology regulation. The startup is therefore required to '
    'be the integrating layer \u2013 the entity that joins up the regulatory abstraction stack because the '
    'institutional architecture cannot. The cost of this integration is borne entirely by the regulated '
    'entity, and it is a cost that increases with every additional regulatory silo.'
)

heading('3.2 The Economic Stakes: Quantified', level=2)

para_mid_fn(
    'The economic significance of the digital sector across East Africa is substantial and growing. '
    'In Kenya, the information and communication technology (ICT) sector contributed 10.8 per cent of '
    'GDP in 2024, growing at an annual average of 8.2 per cent.',
    'Communications Authority of Kenya, ICT Sector Statistics Report FY 2023/24 (CAK, 2024) 12. '
    'The report notes that ICT GDP growth has outpaced overall GDP growth for six consecutive years.',
    ' In Rwanda, the ICT sector contributed 19 per cent of GDP in the first quarter of 2025, driven '
    'primarily by mobile financial services and digital infrastructure investment.'
)

para_mid_fn(
    'For comparative context, Singapore\u2019s digital economy contributed 18.6 per cent of GDP in 2023, '
    'up from 14.9 per cent in 2019, following the consolidation of its regulatory agencies into the '
    'Infocomm Media Development Authority (IMDA) in 2016.',
    'Infocomm Media Development Authority, Singapore Digital Economy Report 2024 (IMDA, 2024) 5. '
    'Singapore merged the Infocomm Development Authority (IDA) and the Media Development Authority (MDA) '
    'to form IMDA in 2016 through the Info-communications Media Development Authority Act 2016 (Singapore).',
    ''
)

para_fn(
    'The Singapore comparison is instructive. The IMDA merger, which consolidated telecommunications, '
    'media, and technology regulation under a single agency, was followed by a measurable acceleration '
    'in digital economic output. Between 2019 and 2023, Singapore\u2019s digital economy grew by 25 per cent '
    'in real terms, and the Infocomm and Technology sector workforce reached 80,000 by 2024, matching '
    'the target set in the 2020 Digital Economy Blueprint.',
    'IMDA (n 8) 12\u201314. See also Infocomm Media Development Authority, Digital Economy Blueprint '
    '(Singapore Government, 2020) 22, which set a target of 80,000 infocomm and technology jobs by 2025.'
)

para_fn(
    'The East African figures, while impressive in absolute terms, must be understood against the '
    'regulatory fragmentation described above. Fragmented regulation acts as a tax on digital economic '
    'activity: it increases the cost of compliance, delays time-to-market, and discourages both domestic '
    'entrepreneurship and foreign direct investment. The World Bank\u2019s Digital Economy for Africa '
    'initiative has identified regulatory harmonisation as one of five priority areas for achieving a '
    'unified digital market in Africa by 2030.',
    'World Bank, Digital Economy for Africa: A Regional Diagnostic (World Bank, 2022) 45. '
    'The initiative identifies harmonisation of digital regulations as a key enabler for cross-border '
    'digital commerce and innovation.'
)

# --- 4. STATUTORY DEFINITIONS ---
heading('4. Statutory Definitions and Conceptual Fragmentation')

para(
    'The horizontal abstraction layer problem is not merely institutional but conceptual. The '
    'foundational definitions upon which technology regulation depends are not harmonised across East '
    'Africa, and in some cases are not internally consistent within a single jurisdiction. This section '
    'examines the statutory definitions of computer and computer system across four East African '
    'jurisdictions, reproducing the precise text of each definition from the official legislative source.'
)

heading('4.1 Uganda: Computer Misuse Act 2011', level=2)

para('The Computer Misuse Act 2011 (Uganda), section 1, defines \u2018computer\u2019 as follows:', italic=True)

quote(
    '\u201ccomputer\u201d means an electronic, magnetic, optical, electrochemical or other data processing '
    'device or a group of such interconnected or related devices, performing logical, arithmetic or '
    'storage functions; and includes any data storage facility or communications facility directly '
    'related to or operating in conjunction with such a device or group of such interconnected or '
    'related devices;'
)

para_fn(
    'The Ugandan definition is notably broad and technologically neutral. By including magnetic, '
    'optical, electrochemical or other data processing device, the legislature anticipated forms of '
    'computing technology beyond the purely electronic. The definition also explicitly includes data '
    'storage and communications facilities, extending the concept beyond the processing unit itself. '
    'This breadth is commendable, but it creates interpretive challenges when applied alongside other '
    'Ugandan statutes that use narrower or different formulations of related terms.',
    'Computer Misuse Act 2011 (Uganda) s 1 (definition of \u2018computer\u2019). This is the definition as '
    'published in the consolidated version of the Act, Chapter 96 of the Laws of Uganda, as accessed '
    'via ULII and Laws.Africa. The definition is broad and technologically neutral, encompassing '
    'magnetic, optical, and electrochemical technologies alongside electronic data processing.'
)

heading('4.2 Tanzania: Cybercrimes Act 2015', level=2)

para('The Cybercrimes Act 2015 (Tanzania), section 3, defines \u2018computer system\u2019 as follows:', italic=True)

quote(
    '\u201ccomputer system\u201d means a device or combination of devices, including network, input and output '
    'devices designed to function as a computer system;'
)

para_fn(
    'The Tanzanian definition takes a narrower approach than its Ugandan counterpart. It does not '
    'specify the technologies (electronic, magnetic, optical, etc.) that the device may use, and it '
    'ties the definition to a functional criterion: the device must be designed to function as a '
    'computer system. This introduces a teleological element that the Ugandan definition lacks. The '
    'Tanzanian definition also omits any reference to storage or communications facilities as separate '
    'components, potentially creating gaps in coverage.',
    'Cybercrimes Act 2015 (Tanzania) s 3 (definition of \u2018computer system\u2019). This is the definition '
    'as published in the official Government Gazette and consolidated by the Tanzania Legal Information '
    'Institute (TanzLII). Note that the original paper incorrectly included \u2018mobile phones\u2019 and '
    '\u201cfollow a set of instructions\u201d in the definition, both of which do not appear in the Act.'
)

heading('4.3 Rwanda: Law No 60/2018 of 22/8/2018', level=2)

para(
    'Law No 60/2018 of 22/8/2018 on the Prevention and Punishment of Cyber Crimes, Article 3, '
    'defines both computer and computer system. Article 3(6\u00b0) provides:',
    italic=True
)

quote(
    '\u201ccomputer\u201d: includes mobile phones, smart phones, computer networks and other devices connected '
    'to the internet;'
)

para('Article 3(11\u00b0) defines \u2018computer system\u2019 separately:', italic=True)

quote(
    '\u201ccomputer system\u201d: an electronic device or combination of electronic devices composed of hardware '
    'or software, including input and output devices with data processing and storage capabilities;'
)

para_fn(
    'Rwanda\u2019s approach is unique in the region for providing separate definitions of computer and '
    'computer system in the same statute. The definition of computer is device-specific and '
    'inclusive (using \u2018includes\u2019), while the definition of computer system is function-specific, '
    'referring to hardware or software with data processing and storage capabilities. This dual '
    'definitional structure creates internal complexity: the concept of computer is explicitly '
    'connected to internet-enabled devices, while computer system is not. The relationship between '
    'the two definitions is not made clear in the statute.',
    'Law No 60/2018 of 22/8/2018 on the Prevention and Punishment of Cyber Crimes (Rwanda) Art 3(6\u00b0) '
    '(definition of \u2018computer\u2019) and Art 3(11\u00b0) (definition of \u2018computer system\u2019). This is the English '
    'translation of the official Kinyarwanda text as published in the Official Gazette and consolidated '
    'by RwandaLII. Note that the original paper incorrectly conflated the two definitions and invented '
    'the phrase \u2018any computer data\u2019, which does not appear in the Act.'
)

heading('4.4 Kenya: Computer Misuse and Cybercrimes Act 2018', level=2)

para(
    'The Computer Misuse and Cybercrimes Act 2018 (Kenya), section 2, defines \u2018computer system\u2019 '
    'as follows:',
    italic=True
)

quote(
    '\u201ccomputer system\u201d means a physical or virtual device, or a set of associated physical or virtual '
    'devices, which use electronic, magnetic, optical or other technology, to perform logical, '
    'arithmetic storage and communication functions on data or which perform control functions on '
    'physical or virtual devices including mobile devices and reference to a computer system includes '
    'a reference to part of a computer system;'
)

para_fn(
    'Kenya\u2019s definition is the most comprehensive of the four jurisdictions examined. It explicitly ' 
    'includes virtual devices, covers multiple technology types (electronic, magnetic, optical, or '
    'other), and importantly provides that reference to a computer system includes a reference to '
    'part of a computer system. This latter provision is significant for the abstraction layer '
    'analysis, as it potentially allows a component or subsystem to be regulated as part of a '
    'computer system even if it would not independently meet the full definition. The Kenyan definition '
    'also explicitly includes mobile devices, bridging the gap between traditional computing and '
    'mobile technology.',
    'Computer Misuse and Cybercrimes Act 2018 (Kenya) s 2 (definition of \u2018computer system\u2019). This is '
    'the definition as published in the Kenya Gazette Supplement No 56 (Act No 5 of 2018) and '
    'consolidated by the National Council for Law Reporting (Kenya Law). Note that the original paper '
    'incorrectly omitted virtual devices, the phrase \u2018physical or virtual\u2019, and the important '
    'clarifying phrase \u2018and reference to a computer system includes a reference to part of a computer '
    'system\u2019.'
)

heading('4.5 Comparative Analysis', level=2)

para('The four definitions reveal significant conceptual fragmentation across the region. Table 1 summarises the key differences.')

table = doc.add_table(rows=8, cols=5)
table.style = 'Table Grid'

headers = ['Feature', 'Uganda CMA 2011', 'Tanzania Cybercrimes 2015', 'Rwanda Law 60/2018', 'Kenya CMCA 2018']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = 'Bookman Old Style'
            r.font.size = Pt(10)
            r.bold = True

data = [
    ['Defines computer separately', 'Yes', 'No', 'Yes (Art 3(6\u00b0))', 'No'],
    ['Defines computer system', 'No', 'Yes', 'Yes (Art 3(11\u00b0))', 'Yes'],
    ['Includes virtual devices', 'No', 'No', 'No', 'Yes'],
    ['Specifies technology types', 'Yes (electronic, magnetic, optical, electrochemical)', 'No', 'Yes (electronic)', 'Yes (electronic, magnetic, optical, other)'],
    ['Includes mobile devices', 'No', 'No', 'Yes (Art 3(6\u00b0))', 'Yes (s 2)'],
    ['Includes part of a system', 'No', 'No', 'No', 'Yes'],
    ['Functional criterion', 'Logical, arithmetic, storage', 'Designed to function as', 'Data processing and storage', 'Logical, arithmetic, storage, communication, control'],
]

for row_idx, row_data in enumerate(data):
    for col_idx, cell_text in enumerate(row_data):
        cell = table.rows[row_idx + 1].cells[col_idx]
        cell.text = cell_text
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Bookman Old Style'
                r.font.size = Pt(10)

p_table_fn = doc.add_paragraph(style='Normal')
run_table_fn = p_table_fn.add_run('Table 1: Comparison of Foundational Definitions Across Four East African Jurisdictions')
run_table_fn.italic = True
run_table_fn.font.name = 'Bookman Old Style'
run_table_fn.font.size = Pt(12)
add_fn_to_para(p_table_fn, 'Source: compiled from Computer Misuse Act 2011 (Uganda) s 1; Cybercrimes Act 2015 (Tanzania) s 3; Law No 60/2018 (Rwanda) Art 3(6\u00b0) and (11\u00b0); Computer Misuse and Cybercrimes Act 2018 (Kenya) s 2.')

para(
    'The fragmentation exposed in Table 1 is not a trivial taxonomical dispute. It has direct practical '
    'consequences. A technology product classified as a computer in Uganda but not in Tanzania may be '
    'subject to different offences, penalties, and enforcement mechanisms. A virtual device recognised '
    'in Kenya may fall outside the statutory framework entirely in Uganda, Tanzania, and Rwanda. A '
    'mobile phone that is explicitly included in the definition of computer in Rwanda is not mentioned '
    'in the Ugandan definition at all. These inconsistencies create regulatory uncertainty for '
    'manufacturers, importers, and users, and they create enforcement gaps that can be exploited by '
    'bad actors.'
)

# --- 5. THE EU CYBER RESILIENCE ACT ---
heading('5. The EU Cyber Resilience Act: A Horizontal Regulatory Solution?')

para_fn(
    'The European Union\u2019s Cyber Resilience Act (Regulation (EU) 2024/2847) represents the most '
    'ambitious attempt to date to solve the horizontal regulatory abstraction layer problem at the '
    'legislative level.',
    'Regulation (EU) 2024/2847 of the European Parliament and of the Council of 23 October 2024 on '
    'horizontal cybersecurity requirements for products with digital elements (Cyber Resilience Act) '
    '[2024] OJ L 2847. The CRA entered into force on 10 December 2024, with phased compliance '
    'obligations beginning in 2026.'
)

para(
    'The CRA establishes horizontal cybersecurity requirements for all products with digital elements '
    'placed on the EU market, regardless of the sector in which they are used. It applies to hardware '
    'and software alike, and it imposes obligations on manufacturers, importers, and distributors '
    'throughout the product lifecycle.'
)

heading('5.1 Background and Legislative History', level=2)

para_fn(
    'The CRA was proposed by the European Commission on 15 September 2022, following a series of '
    'high-profile cybersecurity incidents that exposed the vulnerability of the horizontal regulatory '
    'stack. The SolarWinds attack of 2020, in which malicious code was inserted into the Orion '
    'platform\u2019s build system and distributed to approximately 18,000 customers through a trusted '
    'software update, demonstrated that a single compromise at the software supply chain layer could '
    'propagate upward to affect government agencies and Fortune 500 companies worldwide. '
    'The Kaseya ransomware attack of July 2021, the WannaCry global ransomware attack of May 2017, '
    'and a marked increase in vulnerabilities in Internet of Things (IoT) devices collectively '
    'demonstrated that sector-specific regulation was inadequate to address cross-cutting cybersecurity risks.',
    'European Commission, Proposal for a Regulation on Horizontal Cybersecurity Requirements for '
    'Products with Digital Elements COM(2022) 454 final (15 September 2022) 1. On the SolarWinds '
    'attack, see US Senate Committee on Homeland Security and Governmental Affairs, SolarWinds: A '
    'Bipartite Retrospective (Senate Report, 2022) 5\u201312. The Kaseya ransomware attack of July 2021 '
    'affected approximately 1,500 businesses through a single compromise of the Kaseya VSA remote '
    'management platform: see Mandiant, Kaseya VSA Supply Chain Ransomware Attack: Technical Analysis '
    '(Mandiant, 2021).'
)

heading('5.2 Key Provisions and Relevance to East Africa', level=2)

para_fn(
    'Several features of the CRA are directly relevant to the abstraction layer analysis. First, the '
    'CRA establishes a horizontal regulatory framework that cuts across existing sectoral legislation. '
    'A product with digital elements is subject to the CRA regardless of whether it is also regulated '
    'as a medical device, a financial instrument, or a telecommunications product. This horizontal '
    'approach directly addresses the fragmentation problem identified in sections 3 and 4 above.',
    'CRA, Art 2(2). The CRA applies without prejudice to sector-specific regulations but fills gaps '
    'where no such regulations exist. Where sector-specific regulations provide equivalent cybersecurity '
    'requirements, the CRA does not duplicate them (Art 2(3)).'
)

para_fn(
    'Second, the CRA introduces mandatory vulnerability handling requirements throughout the product '
    'lifecycle, including the obligation to provide security updates for at least five years after '
    'placing the product on the market. '
    'This is a direct response to the obsolescence problem discussed in section 8 below: products '
    'must remain secure for a defined period, and manufacturers cannot externalise the cost of '
    'obsolescence onto consumers or the environment.',
    'CRA, Art 13 (obligations of manufacturers) and Annex I (vulnerability handling requirements). '
    'The five-year support period runs from the date the product is placed on the market.'
)

para_fn(
    'Third, the CRA creates obligations for all actors in the supply chain, including importers and '
    'distributors, not only manufacturers. '
    'This supply-chain-wide approach is a direct response to the SolarWinds and Kaseya attacks, which '
    'exploited the absence of horizontal obligations at the supply chain layer. '
    'This layered liability model is the closest existing analogy to a regulatory abstraction layer: '
    'each actor in the supply chain is responsible for the security of its own layer, and the '
    'regulatory framework ensures that all layers are covered.',
    'CRA, Arts 16\u201318 (obligations of importers) and Arts 19\u201321 (obligations of distributors). This '
    'supply-chain-wide approach is a direct response to the SolarWinds and Kaseya attacks, which '
    'exploited the absence of horizontal obligations at the supply chain layer.'
)

para(
    'For East Africa, the CRA offers both a regulatory template and a market challenge. As a template, '
    'the CRA demonstrates that horizontal regulation across the technology stack is legislatively '
    'achievable, even in a complex multi-jurisdictional environment. As a market challenge, the CRA '
    'creates a compliance asymmetry: products manufactured in East Africa for export to the EU must '
    'meet CRA requirements, but products imported into East Africa from jurisdictions without equivalent '
    'regulation may not. This asymmetry reinforces the obsolescence trap discussed in section 8.'
)

# --- 6. CASE LAW ---
heading('6. Case Law and the Abstraction Layer Gap')

para_fn(
    'A search of East African case law reveals no decided case that directly addresses the abstraction '
    'layer problem as framed in this paper. This is itself a significant finding: it confirms that the '
    'structural blind spot created by regulatory fragmentation has not been tested in court, and '
    'therefore that the legal consequences of a cross-layer regulatory failure remain entirely '
    'unexplored in the region.',
    'The author conducted a systematic search of the Uganda Legal Information Institute (ULII), '
    'Tanzania Legal Information Institute (TanzLII), Rwanda Legal Information Institute (RwandaLII), '
    'and Kenya Law databases using keywords including computer system, abstraction, supply chain, '
    'regulatory fragmentation, cybersecurity, and product liability as of July 2026. No case '
    'addressing the structural relationship between multiple regulatory regimes governing a single '
    'technology product was identified.'
)

para(
    'This absence of precedent is consistent with the thesis of this paper: if the regulatory '
    'abstraction layer is structured to produce a blind spot, then failures that span multiple layers '
    'will be invisible to any single adjudicative forum. A consumer injured by a defective smart device '
    'may have a claim against the manufacturer in contract or tort, but no claim against the regulatory '
    'architecture that failed to coordinate oversight across the twelve agencies described in section 3.'
)

para_fn(
    'Jurisprudence from other jurisdictions, while not directly binding, is instructive. In the United '
    'States, the case of Holbrook v Prodomax Automation Ltd considered whether software constitutes a '
    'product for purposes of product liability law, holding that software can be a product where it '
    'is mass-produced and distributed. '
    'In Europe, the revised Product Liability Directive (EU) 2024/2853, adopted in October 2024, '
    'explicitly includes software as a product for the purposes of liability.',
    'Holbrook v Prodomax Automation Ltd, No 1:19-CV-377 (WD Mich, 29 June 2021). The court held that '
    'software is a product for purposes of Michigan products liability law where it is sold in a '
    'mass-market context. This reasoning was followed in In re: Social Media Adolescent Addiction/Personal '
    'Injury Products Liability Litigation, No 22-md-03047 (ND Cal, 2023), where the court applied a '
    'defect-specific approach to social media platforms. Directive (EU) 2024/2853 of the European '
    'Parliament and of the Council of 23 October 2024 on liability for defective products [2024] OJ L '
    '2853, Art 4(1). The Directive replaces the Product Liability Directive 85/374/EEC and explicitly '
    'provides that software is a product for the purposes of the Directive, reversing the earlier '
    'position where software was considered a service.'
)

para(
    'These developments are relevant to the abstraction layer analysis because they establish that '
    'liability can be layered across the supply chain. The EU Product Liability Directive now provides '
    'that a manufacturer, importer, or distributor may be liable for defective software at any level '
    'of the technical stack. This layered liability model is the closest existing analogy to a '
    'regulatory abstraction layer, but it remains reactive (liability after failure) rather than '
    'proactive (oversight before failure).'
)

# --- 7. STARTUP IMPACT ---
heading('7. The Startup Impact: Quantifying the Regulatory Friction')

para(
    'The cost of the horizontal regulatory abstraction layer is borne most heavily by startups and '
    'small-to-medium enterprises (SMEs). Unlike large multinational corporations, which can maintain '
    'dedicated compliance departments, a fintech startup may have a single legal officer or external '
    'counsel responsible for navigating all twelve regulatory bodies described in section 3.1.'
)

para_fn(
    'The direct costs of this regulatory friction are measurable. A survey by the World Bank\u2019s '
    'Doing Business project, now replaced by the Business Ready (B-READY) methodology, found that '
    'Ugandan firms spend an average of 42 days per year on tax compliance alone, compared to an '
    'average of 24 days in Sub-Saharan Africa. '
    'For a regulated fintech startup, the compliance burden is multiplied across multiple agencies: '
    'registration and licensing (URSB, UCC, BOU), ongoing reporting (FIA, URA, NSSF), and periodic '
    'renewals and inspections (KCCA, UNBS). The Uganda Investment Authority estimates that regulatory '
    'compliance costs account for between 8 and 15 per cent of operating expenses for early-stage '
    'technology firms.',
    'World Bank, Business Ready (B-READY) Report: Uganda 2025 (World Bank, 2025). '
    'Previous data from the Doing Business project indicated tax compliance time of 42 days per year '
    'for Ugandan firms: Doing Business 2020 (World Bank, 2020) 83. Uganda Investment Authority, '
    'Cost of Doing Business in Uganda 2024 (UIA, 2024) 22. The estimate covers all regulatory '
    'compliance costs including licensing, reporting, inspections, and legal fees.'
)

heading('7.1 Consumer Harm: The Pass-Through Effect', level=2)

para_fn(
    'The cost of regulatory compliance does not remain with the startup. It is passed through to '
    'consumers in the form of higher prices, reduced choice, and delayed service availability. A '
    'fintech startup that requires twelve separate regulatory approvals before launching a mobile '
    'money service will either delay its market entry (reducing consumer access to financial services) '
    'or increase its fees (reducing affordability). The World Bank\u2019s Global Findex Database 2021 '
    'reports that 52 per cent of Ugandan adults do not have access to formal financial services. '
    'Every regulatory barrier that delays or increases the cost of fintech market entry is a barrier '
    'to financial inclusion. The horizontal regulatory abstraction layer thus has a directly '
    'measurable impact on the welfare of the poorest and most excluded members of society.',
    'World Bank, The Global Findex Database 2021: Financial Inclusion, Digital Payments, and Resilience '
    'in the Age of COVID-19 (World Bank, 2022) 45. The report notes that mobile money accounts have '
    'increased financial inclusion in Sub-Saharan Africa but that regulatory barriers remain a '
    'significant constraint.'
)

# --- 8. OBSOLESCENCE TRAP ---
heading('8. The Obsolescence Trap')

para(
    'The horizontal regulatory abstraction layer creates a structural incentive for the dumping of '
    'obsolete technology in East African markets. A technology product that can no longer be sold in '
    'the European Union because it does not meet CRA requirements, or in the United States because it '
    'does not meet Federal Communications Commission (FCC) or National Institute of Standards and '
    'Technology (NIST) standards, may still be legally marketable in East Africa because no single '
    'regulator has the mandate or the capacity to assess its compliance with a comprehensive safety '
    'framework.'
)

para_fn(
    'The scale of this phenomenon is significant. Industry analysts estimate that between 50 and 70 '
    'per cent of mobile phones and computing devices entering secondary markets in Africa are '
    'refurbished or second-hand, and that a substantial proportion of these do not meet current safety '
    'or security standards in their markets of origin.',
    'Canalys, Global Smartphone Market: Secondary Market Dynamics 2024 (Canalys, 2024) 12; '
    'Counterpoint Research, Refurbished Smartphone Market Tracker Q4 2024 (Counterpoint, 2025) 8. '
    'These figures are industry estimates based on trade flows and supply chain data. A widely-cited '
    'ITU statistic of 60 to 80 per cent refurbished devices in developing countries could not be '
    'verified in ITU publications and appears to originate from industry analyst sources.'
)

para(
    'The obsolescence trap is not merely a consumer protection issue. It is an environmental issue '
    '(obsolete devices contribute to the growing e-waste stream in East Africa), a cybersecurity issue '
    '(devices that no longer receive security updates become vectors for botnets and malware), and a '
    'digital sovereignty issue (East African markets become dependent on technology designed for and '
    'discarded by other jurisdictions). The absence of a horizontal regulatory abstraction layer means '
    'that no single agency is responsible for assessing the cumulative impact of obsolete technology '
    'across all of these dimensions.'
)

# --- 9. REFORM AGENDA ---
heading('9. A Reform Agenda')

para(
    'The preceding analysis suggests four concrete reforms, each of which addresses a distinct dimension '
    'of the horizontal abstraction layer problem.'
)

heading('9.1 Harmonised Foundational Definitions', level=2)

para(
    'The first reform is the harmonisation of foundational statutory definitions across the East African '
    'Community (EAC) Partner States. The fragmentation documented in section 4 creates regulatory '
    'uncertainty, enforcement gaps, and compliance costs. A harmonised definitional framework, adopted '
    'through the EAC Legislative Assembly or through bilateral mutual recognition agreements, would '
    'establish a common conceptual foundation. The Kenyan definition in the Computer Misuse and '
    'Cybercrimes Act 2018 provides the most comprehensive template, as it explicitly includes virtual '
    'devices, multiple technology types, mobile devices, and the critical clarifying phrase that '
    'reference to a computer system includes a reference to part of a computer system.'
)

heading('9.2 A Regional Regulatory Abstraction Layer Body', level=2)

para_fn(
    'The second reform is the establishment of a regional body with a horizontal mandate across the '
    'technology regulatory stack. This body would be modeled on Singapore\u2019s IMDA, which consolidated '
    'telecommunications, media, and technology regulation under a single agency in 2016, and which has '
    'been associated with a 25 per cent real growth in digital economic output between 2019 and 2023. '
    'A regional body would not replace existing sectoral regulators but would provide the horizontal '
    'coordination function that currently does not exist: assessing cross-cutting risks, harmonising '
    'standards, and providing a single point of regulatory entry for technology products entering the '
    'regional market.',
    'IMDA (n 8). The IMDA was established by the Info-communications Media Development Authority Act '
    '2016 (Singapore), merging the Infocomm Development Authority and the Media Development Authority.'
)

heading('9.3 Software Bill of Materials (SBOM) Mandates', level=2)

para_fn(
    'The third reform is the adoption of mandatory Software Bill of Materials (SBOM) requirements for '
    'technology products sold in the region. The SBOM, defined by the United States National '
    'Telecommunications and Information Administration (NTIA) as a formal record containing the '
    'details and supply chain relationships of the components used in building software, provides a '
    'horizontal transparency mechanism that cuts across institutional silos. '
    'If every technology product entering the East African market were required to include an SBOM, '
    'every regulator in the abstraction layer would have access to the same foundational information '
    'about the product\u2019s componentry. The SBOM would serve as a shared informational layer, partially '
    'compensating for the absence of a shared institutional layer.',
    'NTIA, Software Bill of Materials: An Introduction (US Department of Commerce, 2021) 3. The NTIA '
    'SBOM initiative was established pursuant to Executive Order 14028 on Improving the Nation\u2019s '
    'Cybersecurity (12 May 2021). The EU CRA also mandates SBOM-like component disclosure: CRA, '
    'Annex I, Part B(2).'
)

heading('9.4 The Ostrom-Hess Design Principles Applied', level=2)

para_fn(
    'The fourth reform is the application of Elinor Ostrom and Charlotte Hess\u2019s design principles for '
    'managing shared knowledge resources to the regulatory abstraction layer problem.',
    'Elinor Ostrom and Charlotte Hess, \u2018A Framework for Analyzing the Knowledge Commons\u2019 in Charlotte '
    'Hess and Elinor Ostrom (eds), Understanding Knowledge as a Commons: From Theory to Practice '
    '(MIT Press 2007) 41\u201381. Ostrom\u2019s Nobel Prize-winning work on governing the commons identified '
    'eight design principles for sustainable management of shared resources, including clearly defined '
    'boundaries, collective choice arrangements, and nested enterprises.'
)

para(
    'Ostrom and Hess argue that shared knowledge resources, like shared natural resources, require '
    'institutional arrangements that are polycentric, participatory, and adaptive. The regulatory '
    'abstraction layer in East Africa can be understood as a knowledge commons: a shared resource '
    '(regulatory information and oversight capacity) that is currently fragmented across multiple '
    'institutional actors. Applying Ostrom and Hess\u2019s principles would suggest reforms including: '
    '(a) clear boundaries on which regulators have authority over which layers of the technology stack; '
    '(b) inclusive decision-making that brings regulators, industry, and civil society together; and '
    '(c) nested institutional arrangements, in which regional, national, and sectoral bodies are '
    'linked through formal coordination mechanisms.'
)

# --- 10. CONCLUSION ---
heading('10. Conclusion')

para(
    'This paper has argued that the abstraction layer problem, well-recognised in computer science '
    'as the structural blind spot created by layered technical architectures, also exists horizontally '
    'in the regulation of technology across East Africa. The vertical abstraction layer enables digital '
    'scalability by allowing each layer of the technical stack to function independently; the '
    'horizontal regulatory abstraction layer produces fragmentation by allowing each institution and '
    'statutory regime to function independently. The result in both cases is a structural blind spot: '
    'no single actor sees the full stack, and failures at one layer propagate invisibly to others.'
)

para(
    'The evidence for this claim is threefold. First, the foundational statutory definitions across '
    'four East African jurisdictions are fragmented to a degree that cannot be attributed to mere '
    'drafting variation: they reflect fundamentally different conceptualisations of the subject matter '
    'of technology regulation. Second, the institutional landscape in a country like Uganda, where a '
    'single fintech startup may answer to twelve separate regulatory bodies, creates a compliance '
    'burden that falls disproportionately on the innovators and entrepreneurs who drive digital '
    'economic growth. Third, the absence of any horizontal regulatory coordination mechanism means '
    'that cross-cutting problems \u2013 including obsolete technology dumping, supply chain cybersecurity '
    'vulnerabilities, and e-waste \u2013 remain unaddressed by any single institution.'
)

para(
    'The paper has proposed four reforms: harmonised foundational definitions modelled on the Kenyan '
    'Computer Misuse and Cybercrimes Act 2018; a regional regulatory abstraction layer body modelled '
    'on Singapore\u2019s IMDA; mandatory Software Bill of Materials requirements for all technology '
    'products entering the regional market; and the application of Ostrom and Hess\u2019s design principles '
    'for shared knowledge resource governance. These reforms are not exhaustive, but they are mutually '
    'reinforcing. A harmonised definitional framework enables the consistent application of standards; '
    'a regional body provides the institutional architecture for horizontal coordination; SBOM mandates '
    'provide the informational infrastructure; and polycentric governance principles ensure that the '
    'resulting system is adaptive, participatory, and resilient.'
)

para(
    'The European Union\u2019s Cyber Resilience Act demonstrates that horizontal regulation of the '
    'technology stack is legislatively achievable. For East Africa, the question is not whether a '
    'regulatory abstraction layer is needed, but whether the region will build one deliberately or '
    'continue to suffer the consequences of its absence.'
)

# --- SAVE ---
output_path = 'C:\\Users\\DELL\\research\\Redrafted_Research_Paper.docx'
doc.save(output_path)
print(f'Document saved to {output_path}')

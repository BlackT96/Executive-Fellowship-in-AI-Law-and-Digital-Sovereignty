from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

with open(r'C:\Users\DELL\research\Agents\writing\article_draft_ughub_whatsapp.md', 'r', encoding='utf-8') as f:
    text = f.read()

lines = text.split('\n')

for line in lines:
    if line.startswith('# '):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line[2:])
        run.bold = True
        run.font.size = Pt(20)
    elif line.startswith('## '):
        p = doc.add_paragraph()
        run = p.add_run(line[3:])
        run.bold = True
        run.font.size = Pt(14)
    elif line.startswith('### '):
        p = doc.add_paragraph()
        run = p.add_run(line[4:])
        run.bold = True
        run.font.size = Pt(12)
    elif line.startswith('---'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('\u2014' * 40)
        run.font.color.rgb = RGBColor(150, 150, 150)
    elif line.startswith('*') and line.endswith('*') and len(line) > 2:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line[1:-1])
        run.italic = True
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(80, 80, 80)
    elif line.strip() == '':
        pass
    else:
        parts = re.split(r'(\*\*.*?\*\*)', line)
        p = doc.add_paragraph()
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
                run.font.name = 'Calibri'
                run.font.size = Pt(11)
            else:
                run = p.add_run(part)
                run.font.name = 'Calibri'
                run.font.size = Pt(11)

doc.save(r'C:\Users\DELL\research\Agents\writing\UGHUB_WhatsApp_Article.docx')
print("Done")

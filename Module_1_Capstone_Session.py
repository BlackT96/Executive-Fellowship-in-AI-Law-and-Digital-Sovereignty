import docx
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = docx.Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
pf = style.paragraph_format
pf.space_after = Pt(6)

def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def para(text, bold=False, italic=False, size=None, color=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold: run.bold = True
    if italic: run.italic = True
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def shaded_cell(cell, text, bold=False, color=None, shading=None):
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    if bold: run.bold = True
    if color: run.font.color.rgb = RGBColor(*color)
    if shading:
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), shading)
        shading_elm.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(shading_elm)

def set_col_widths(table, widths):
    for row in table.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                row.cells[i].width = Cm(w)

# ════════════════════════════════════════════════════════════
# TITLE PAGE
# ════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EXECUTIVE FELLOWSHIP IN\nAI LAW & DIGITAL SOVEREIGNTY")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0, 40, 80)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Module 1 — Capstone Session")
run.bold = True
run.font.size = Pt(16)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Digital Technology Fundamentals\nFull-Stack Legal Map, Reference Tables & Integrated Problem")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Duration: 90–120 minutes  |  Format: Open-book, group discussion")
run.font.size = Pt(10)
run.italic = True

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# PART 1: THE FULL-STACK LEGAL MAP
# ════════════════════════════════════════════════════════════
heading("PART 1 — The Full-Stack Legal Map", level=1)

para("The Full-Stack Legal Map is a single visual framework that connects every technical concept from Module 1 to its legal significance. It answers the question: 'When something goes wrong in a computer system, which law applies, and who is liable?'", italic=True)

doc.add_paragraph()

# ── The Stack Diagram ──
heading("1.1  The Four-Layer Stack", level=2)

para("Every digital system — from a mobile money app to a hospital records platform — can be understood as four stacked layers. A failure at any layer triggers different legal questions:", italic=True, size=10)

doc.add_paragraph()

# ── Build the stack as a table ──
stack_data = [
    ["", "LAYER 4\nAPI & CLOUD", "What connects apps\nto data and infrastructure", "Cloud Agreement • API Audit Log\nETA s.29 (service provider liability)\nDPPA s.20 (data processor obligations)\nContract law (SLA, warranties, indemnities)", "Global: frictionless cloud\naccess, standard SLAs\n\nUganda: cross-border data\nrestrictions, weak bargaining\npower vs foreign providers"],
    ["", "LAYER 3\nDATABASE", "Where data is\nstored and queried", "ETA s.8(4) & s.8(5) (evidential weight,\npresumptions for electronic records)\nETA s.9 (accessibility & retention)\nDPPA s.18 (data minimisation)\nLimitation Act Cap. 80 (retention periods)", "Global: ACID compliance =\nreliable evidence\n\nUganda: grey-market hardware,\nopen-source DBs without\nSLAs affect reliability"],
    ["", "LAYER 2\nNETWORK", "How data travels\nbetween devices", "RICA (lawful interception)\nCMA s.12, s.14, s.15 (unauthorised\naccess, intent, modification)\nETA s.7 (authenticity of data message)\nJurisdiction (where data passed)", "Global: assume unmetered\nhigh-speed, stable routing\n\nUganda: erratic infra,\nmobile-first, data passing\nthrough multiple carriers"],
    ["", "LAYER 1\nHARDWARE & OS", "The physical machine\nand its brain", "Sale of Goods Act Cap. 82 (fitness\nfor purpose, implied terms)\nCMA s.12 (unauthorised access)\nAbstraction layer analysis\n(product liability)", "Global: standardised\nprocurement, clear OEM\nwarranties\n\nUganda: refurbished/\ngrey-market hardware,\nprivity of contract issues"],
]

t = doc.add_table(rows=5, cols=5)
t.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
headers = ["", "LAYER", "PLAIN ENGLISH\n(What it does)", "LEGAL FRAMEWORK\n(What law applies)", "FOUNDATION-TO-TUNE\n(Global vs Uganda)"]
for i, h in enumerate(headers):
    shaded_cell(t.rows[0].cells[i], h, bold=True, color=(255,255,255), shading="1F4E79")
    t.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Layer 1 row
layer_labels = ["⬆︎\nHIGHER\nABSTRACTION\n⬆︎", "", "", "", ""]
layer_colors = ["E8F0FE", "FFF3E0", "E8F5E9", "FCE4EC"]

for idx, (label, data) in enumerate(zip(layer_labels, stack_data)):
    r = t.rows[idx + 1]
    # First cell — layer number
    shaded_cell(r.cells[0], str(4 - idx), bold=True, shading="D6E4F0")
    r.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data cells
    for col in range(1, 5):
        shaded_cell(r.cells[col], data[col - 1], shading=layer_colors[idx] if col > 0 else None)

set_col_widths(t, [1.2, 2.5, 3.5, 5.5, 4.5])

doc.add_paragraph()

para("Key Insight: Most real disputes involve multiple layers. The skill is isolating which layer(s) the evidence points to, then applying the correct legal framework to each.", bold=True, size=11, color=(0, 60, 120))

doc.add_paragraph()

# ── Decision Tree ──
heading("1.2  Failure Diagnosis Decision Tree", level=2)

para("When investigating a technology dispute, trace the failure through this sequence:", italic=True, size=10)

doc.add_paragraph()

decision_tree = [
    ("1. Is there physical damage?", "Hardware failure → Sale of Goods Act / product liability → Layer 1", "If no physical damage, move to 2"),
    ("2. Was there unauthorised access?", "OS/network intrusion → CMA s.12-15 → Layer 1 or 2", "If no intrusion, move to 3"),
    ("3. Is there a record conflict?", "Database inconsistency → ETA s.8(4)-8(5) → Layer 3", "If no conflict, move to 4"),
    ("4. Is there a service failure?", "SLA breach / API error → Contract / ETA s.29 → Layer 4", "If no service failure → systemic or cross-layer issue"),
]

t2 = doc.add_table(rows=5, cols=3)
t2.alignment = WD_TABLE_ALIGNMENT.CENTER

shaded_cell(t2.rows[0].cells[0], "DIAGNOSTIC QUESTION", bold=True, color=(255,255,255), shading="1F4E79")
shaded_cell(t2.rows[0].cells[1], "IF YES →", bold=True, color=(255,255,255), shading="1F4E79")
shaded_cell(t2.rows[0].cells[2], "IF NO →", bold=True, color=(255,255,255), shading="1F4E79")

for i, (q, yes, no) in enumerate(decision_tree):
    shaded_cell(t2.rows[i+1].cells[0], q, bold=True, shading="F5F5F5")
    shaded_cell(t2.rows[i+1].cells[1], yes, shading="E8F5E9")
    shaded_cell(t2.rows[i+1].cells[2], no, shading="FFF3E0")

set_col_widths(t2, [5, 6.5, 6.5])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# PART 2: REFERENCE TABLES
# ════════════════════════════════════════════════════════════
heading("PART 2 — Reference Tables", level=1)

para("These tables consolidate every key concept from Weeks 1–4 by legal question type. Use them during the capstone exercise.", italic=True)

doc.add_paragraph()

# ── Table 1: Key Concepts x Legal Significance ──
heading("2.1  Technical Concept → Legal Significance", level=2)

concepts = [
    ("Abstraction layer", "A system failure originates at one specific layer. Liability follows the layer.", "Week 1"),
    ("Kernel vs User Mode", "Kernel mode = full system access. User mode = restricted. Malware that escalates to kernel = CMA s.12 (unauthorised modification).", "Week 1"),
    ("Volatile vs Non-Volatile Memory", "RAM, cache, registers lose data on power-off. Forensic preservation must happen before shutdown.", "Week 1"),
    ("TCP/IP 5-Layer Model", "Each layer creates a different jurisdictional hook. IP header = geolocation. Application data = content jurisdiction.", "Week 2"),
    ("Encapsulation", "Each layer wraps data from above. The outer layers (network, transport) determine routing; inner layers (application) determine content.", "Week 2"),
    ("DNS Resolution", "DNS reveals which domain names resolve to which IPs. Critical for identifying hosting jurisdiction and content distribution.", "Week 2"),
    ("ACID vs BASE", "ACID (PostgreSQL) = immediate consistency, reliable evidence. BASE (Cassandra) = eventual consistency, unreliable for point-in-time queries.", "Week 3"),
    ("ETL Pipeline", "Extract, Transform, Load — data moves between systems. Each step can introduce error. Audit trail is key to evidence weight.", "Week 3"),
    ("Serverless Architecture", "No dedicated server = no 'proper working of computer' presumption under ETA s.8(5)(a). Affects evidential weight.", "Week 3"),
    ("REST API & HTTP Verbs", "GET = read, POST = create, PUT = update, DELETE = remove. Audit logs show exactly what operation was performed.", "Week 4"),
    ("IDOR (Broken Access Control)", "User can access another user's data by changing an ID in the URL. Fix: authorisation check per resource.", "Week 4"),
    ("Cloud Service Models", "IaaS = you manage OS/apps. PaaS = you manage apps only. SaaS = you manage nothing. Liability allocation differs per model.", "Week 4"),
    ("SDLC Lifecycle", "Requirements → Design → Dev → Testing → Deployment → Maintenance. Skipping testing = evidence of negligence.", "Week 4"),
    ("SLA & Service Credits", "Sole remedy is often a discretionary credit. Must negotiate security-incident-specific remedies separately.", "Week 4"),
]

t3 = doc.add_table(rows=len(concepts)+1, cols=3)
t3.alignment = WD_TABLE_ALIGNMENT.CENTER

shaded_cell(t3.rows[0].cells[0], "TECHNICAL CONCEPT", bold=True, color=(255,255,255), shading="1F4E79")
shaded_cell(t3.rows[0].cells[1], "LEGAL SIGNIFICANCE", bold=True, color=(255,255,255), shading="1F4E79")
shaded_cell(t3.rows[0].cells[2], "WEEK", bold=True, color=(255,255,255), shading="1F4E79")

for i, (concept, significance, week) in enumerate(concepts):
    shaded_cell(t3.rows[i+1].cells[0], concept, bold=True, shading="F0F4F8")
    shaded_cell(t3.rows[i+1].cells[1], significance)
    shaded_cell(t3.rows[i+1].cells[2], week)
    t3.rows[i+1].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

set_col_widths(t3, [4, 10, 2])

doc.add_paragraph()

# ── Table 2: Statute Quick Reference ──
heading("2.2  Statute Quick Reference", level=2)

statutes = [
    ("Computer Misuse Act (CMA) Cap. 98", "S.12 — Unauthorised access\nS.14 — Access with intent to commit offence\nS.15 — Unauthorised modification", "Criminal liability for\nhacking, intrusion,\nmalware", "Layer 1 (OS)\nLayer 2 (Network)"),
    ("Electronic Transactions Act (ETA) Cap. 99", "S.7 — Authenticity of data messages\nS.8(1)(c) — Admissibility despite not original\nS.8(4) — Factors for evidential weight\nS.8(5) — Presumptions for reliable systems\nS.9 — Accessibility and retention\nS.29 — Service provider immunity", "Admissibility and\nweight of electronic\nevidence", "Layer 3 (Database)\nLayer 4 (Cloud)"),
    ("Data Protection & Privacy Act (DPPA) 2019", "S.18 — Data minimisation & retention\nS.20 — Security safeguards\nSection 19 — Cross-border transfer\nPDPO notification obligation", "Obligations of data\ncontrollers and\nprocessors", "Layer 3 (Database)\nLayer 4 (Cloud)"),
    ("Sale of Goods Act Cap. 82", "S.40-44 — Implied terms (quality, fitness)\nPrivity of contract rules", "Product liability for\ndefective hardware\nor software", "Layer 1 (Hardware)"),
    ("Regulation of Interception of Communications Act (RICA)", "S.3-8 — Lawful interception requirements\nS.12-14 — Unlawful interception offences", "Legal boundaries of\nnetwork surveillance", "Layer 2 (Network)"),
    ("Limitation Act Cap. 80", "6 years for contractual claims\n3 years for tort claims", "Determines minimum\ndata retention periods", "Layer 3 (Database)"),
    ("Reciprocal Enforcement of Judgments Act Cap. 21", "Registration of foreign judgments\nConditions for enforcement", "Enforceability of\nowerseas judgments", "Layer 4 (Cloud)"),
]

t4 = doc.add_table(rows=len(statutes)+1, cols=4)
t4.alignment = WD_TABLE_ALIGNMENT.CENTER

shaded_cell(t4.rows[0].cells[0], "STATUTE", bold=True, color=(255,255,255), shading="1F4E79")
shaded_cell(t4.rows[0].cells[1], "KEY SECTIONS", bold=True, color=(255,255,255), shading="1F4E79")
shaded_cell(t4.rows[0].cells[2], "WHEN TO USE", bold=True, color=(255,255,255), shading="1F4E79")
shaded_cell(t4.rows[0].cells[3], "LAYER", bold=True, color=(255,255,255), shading="1F4E79")

for i, (statute, sections, when, layer) in enumerate(statutes):
    shaded_cell(t4.rows[i+1].cells[0], statute, bold=True, shading="F0F4F8")
    shaded_cell(t4.rows[i+1].cells[1], sections)
    shaded_cell(t4.rows[i+1].cells[2], when)
    shaded_cell(t4.rows[i+1].cells[3], layer)

set_col_widths(t4, [4.5, 5.5, 4, 2.5])

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# PART 3: CAPSTONE PROBLEM
# ════════════════════════════════════════════════════════════
heading("PART 3 — Capstone Problem: The SwiftMove Dispute", level=1)

para("This is an open-book, discussion-based exercise. Use the Full-Stack Legal Map and Reference Tables to guide your analysis.", italic=True)

doc.add_paragraph()

heading("3.1  Fact Pattern", level=2)

para("SwiftMove Uganda Ltd is a Kampala-based logistics company operating a digital delivery platform. The platform consists of:")

bullet("A mobile app for delivery drivers (Android)")
bullet("A cloud backend hosted by DataVault Kenya Ltd (a PaaS provider with servers in Nairobi)")
bullet("A PostgreSQL database storing all delivery records")
bullet("A REST API that drivers' phones call to update delivery status")

doc.add_paragraph()
para("The Incident:", bold=True)

para("On 15 August 2026, a dispute arises over Delivery #48291. The facts are:")

bullet("Driver D-201 (assigned to Delivery #48291) marked the delivery as 'Completed' in the app at 14:32.")
bullet("The customer, Sarah Nakato, says the package never arrived.")
bullet("The PostgreSQL database shows: delivery_status = 'completed', completed_at = '2026-08-15 14:32:04', completed_by = 'D-201'.")
bullet("The API audit log shows a PUT request to /api/v1/deliveries/48291/status at 14:32:04 from D-201's device. The request body contained: {status: 'completed'}.")
bullet("SwiftMove's operations manager swears D-201 was unauthorised to handle this delivery — it was reassigned to Driver D-207 at 13:45 due to D-201 being off-shift, but the reassignment was done by a phone call (no system update).")
bullet("DataVault Kenya Ltd confirms 99.97% uptime for August 2026 and states its infrastructure logs show no anomalies.")
bullet("D-201's phone is a refurbished Tecno phone purchased from a vendor in Owino market. D-201 has since disappeared.")

doc.add_paragraph()
heading("3.2  The Legal Questions", level=2)

para("Advise SwiftMove Uganda Ltd on each of the following. Use the Full-Stack Legal Map to identify which layer(s) each question engages.", bold=True)

doc.add_paragraph()

para("Question 1 — Evidence Reliability (30 mins)", bold=True, size=12, color=(0, 60, 120))
para("Sarah Nakato threatens to sue SwiftMove for the missing package. SwiftMove's only evidence that delivery occurred is the database record and API log.")
bullet("Is the PostgreSQL database record admissible as evidence? (Consider ETA s.8(4), s.8(5))")
bullet("Does the API audit log corroborate or undermine the database record?")
bullet("What factors affect the evidential weight of each record?")
bullet("If this were a criminal matter (theft of the package by D-201), how would your analysis change?")

doc.add_paragraph()
para("Question 2 — Hardware & OS Layer (15 mins)", bold=True, size=12, color=(0, 60, 120))
bullet("D-201's phone is refurbished — does that create any product liability claim against the vendor? (Consider Sale of Goods Act Cap. 82)")
bullet("If D-201's phone had been compromised (malware that auto-sent the PUT request), which abstraction layer failed?")
bullet("Who would be the proper defendant for a hardware-layer failure?")

doc.add_paragraph()
para("Question 3 — Network Layer (15 mins)", bold=True, size=12, color=(0, 60, 120))
bullet("The API request from D-201's phone passed through three mobile carriers to reach DataVault's servers in Nairobi. If the request was intercepted or modified in transit, which law applies?")
bullet("Can SwiftMove determine where along the network path the data was at any given time? (Consider TCP/IP encapsulation)")
bullet("What jurisdictional issues arise from data crossing from Uganda to Kenya?")

doc.add_paragraph()
para("Question 4 — API & Cloud Layer (20 mins)", bold=True, size=12, color=(0, 60, 120))
bullet("SwiftMove's agreement with DataVault Kenya Ltd is governed by English law with exclusive London jurisdiction. DataVault's SLA promises 99.9% uptime — they achieved 99.97%. Does SwiftMove have any remedy?")
bullet("The reassignment of Delivery #48291 was done by phone call, not system update. Is this a process failure (SDLC) or a legal issue?")
bullet("Draft one amendment to the DataVault agreement that would have protected SwiftMove better in this scenario.")

doc.add_paragraph()
para("Question 5 — Integrated Opinion (20 mins)", bold=True, size=12, color=(0, 60, 120))
para("Write a memorandum to SwiftMove's managing director that:")
bullet("Identifies the most likely layer(s) where the failure occurred")
bullet("Assesses the strength of SwiftMove's position in a dispute with Sarah Nakato")
bullet("Recommends three practical steps SwiftMove should take immediately")

doc.add_page_break()

# ════════════════════════════════════════════════════════════
# PART 4: SESSION GUIDE
# ════════════════════════════════════════════════════════════
heading("PART 4 — Session Guide for the Instructor", level=1)

para("This guide walks through how to run the 90–120 minute session. Adjust timing based on group size and pace.", italic=True)

doc.add_paragraph()

heading("4.1  Session Structure", level=2)

session_plan = [
    ("0:00 – 0:15", "Introduce the Full-Stack Legal Map", "Display the 4-layer stack. Walk through each layer: what it does, what law applies, the Foundation-to-Tune difference. Emphasise that most real disputes span multiple layers."),
    ("0:15 – 0:30", "Read the Fact Pattern Together", "Distribute the SwiftMove case. Read aloud. Ask candidates to identify which weeks' concepts appear. (Expected: all four weeks — phone hardware = Wk1, network carriers = Wk2, database = Wk3, API/cloud = Wk4)."),
    ("0:30 – 1:00", "Questions 1 & 2 (Evidence + Hardware)", "Work through together. Key teaching points: ETA s.8(4) factors, s.8(5) presumptions, ACID vs BASE reliability, Sale of Goods implied terms vs privity problem for grey-market goods."),
    ("1:00 – 1:20", "Questions 3 & 4 (Network + Cloud)", "Key teaching points: TCP/IP encapsulation and jurisdiction, SLA analysis, SDLC process failure as evidence of negligence, contract amendment drafting."),
    ("1:20 – 1:50", "Question 5 — Integrated Opinion", "Give candidates 20 minutes to draft a short memorandum. Then compare approaches. Discuss which layer each candidate prioritised."),
    ("1:50 – 2:00", "Debrief & Module 1 Wrap", "Summarise key takeaways. Identify which areas each candidate should revise before Module 2."),
]

t5 = doc.add_table(rows=len(session_plan)+1, cols=3)
t5.alignment = WD_TABLE_ALIGNMENT.CENTER

shaded_cell(t5.rows[0].cells[0], "TIME", bold=True, color=(255,255,255), shading="1F4E79")
shaded_cell(t5.rows[0].cells[1], "ACTIVITY", bold=True, color=(255,255,255), shading="1F4E79")
shaded_cell(t5.rows[0].cells[2], "INSTRUCTOR NOTES", bold=True, color=(255,255,255), shading="1F4E79")

for i, (time, activity, notes) in enumerate(session_plan):
    shaded_cell(t5.rows[i+1].cells[0], time, bold=True, shading="F0F4F8")
    shaded_cell(t5.rows[i+1].cells[1], activity, bold=True)
    shaded_cell(t5.rows[i+1].cells[2], notes)

set_col_widths(t5, [2.5, 4.5, 9.5])

doc.add_paragraph()

heading("4.2  Expected Answers (for Instructor Reference)", level=2)

para("Question 1 — Evidence Reliability:", bold=True, size=11)
bullet("PostgreSQL uses ACID — the record is likely reliable under ETA s.8(4) (manner of generation, reliability of process). The s.8(5) rebuttable presumption of proper working may apply.")
bullet("API audit log corroborates the database record — timestamp and request body match. However, it only proves a request was sent, not that delivery physically occurred.")
bullet("Criminal case: higher standard. Must prove beyond reasonable doubt. D-201's disappearance, refurbished phone, and phone-call reassignment all create reasonable doubt.")
bullet("Key distinction: civil = balance of probabilities, criminal = beyond reasonable doubt.")

para("Question 2 — Hardware & OS:", bold=True, size=11)
bullet("Refurbished phone: Sale of Goods Act Cap. 82 s.40-44 imply terms of satisfactory quality and fitness for purpose. But privity of contract — D-201 bought it, not SwiftMove. SwiftMove has no direct claim against the vendor.")
bullet("If malware: failure at Layer 1 (OS) — kernel/user mode escalation. CMA s.12-15 apply to the attacker. SwiftMove's liability depends on whether they provided the phone or D-201 used his own device.")

para("Question 3 — Network:", bold=True, size=11)
bullet("TCP/IP encapsulation: each layer adds headers. The IP header reveals source/destination. The application data (API request body) is at Layer 5. Interception could occur at any layer.")
bullet("Jurisdiction: data passed through Uganda → Kenya. RICA governs interception in Uganda. Kenya's Data Protection Act governs data in Kenya. Three jurisdictions involved.")

para("Question 4 — API & Cloud:", bold=True, size=11)
bullet("SLA: 99.97% exceeds 99.9%. No remedy under current SLA. The real issue is not uptime but security/authorisation — the current SLA doesn't cover this.")
bullet("Phone-call reassignment: SDLC process failure. The system allowed D-201 to complete a delivery he was not authorised for because the reassignment was never entered into the system. This is a process gap, not a technical bug.")
bullet("Suggested amendment: add clause requiring that access control changes (reassignments, role changes) must be made through the system and cannot be overridden by manual processes.")

para("Question 5 — Integrated Opinion:", bold=True, size=11)
bullet("Likely failure layers: Layer 1 (phone hardware possibly compromised?), Layer 4 (process failure in SDLC — reassignment not reflected in system).")
bullet("Strength of position: weak — SwiftMove has a database record and API log showing completion, but cannot prove physical delivery. The reassignment gap undermines their process.")
bullet("Recommendations: (1) Implement mandatory system-based reassignment — no phone-call overrides. (2) Preserve all logs forensically. (3) Review cloud agreement for security/process obligations.")

doc.add_paragraph()

heading("4.3  Discussion Prompts", level=2)

para("Use these prompts to deepen the session:", italic=True)
doc.add_paragraph()

bullet("If SwiftMove had used Cassandra instead of PostgreSQL, would your evidence analysis change?")
bullet("Should companies be liable for actions their employees take on personal (refurbished) devices?")
bullet("DataVault's SLA says 'as is, as available' — is it fair to hold a Ugandan startup to the same standard as a multinational?")
bullet("The reassignment was done by phone call. Is this a legal issue or a management issue?")

doc.add_paragraph()

# ── Final page ──
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— End of Module 1 Capstone Session —")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(100, 100, 100)

# ── Save ──
path = r'C:\Users\DELL\research\Module 1 - Capstone Session.docx'
doc.save(path)
print(f"Saved to {path}")

# -*- coding: utf-8 -*-
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED = RGBColor(200, 0, 0)
GREEN = RGBColor(0, 100, 0)
BLUE = RGBColor(0, 0, 120)

def add_comment(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = True
    return p

def add_normal(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_mark(doc, label, marks_awarded, marks_total, comment=""):
    p = doc.add_paragraph()
    run = p.add_run(f"[{label}: {marks_awarded}/{marks_total}]")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREEN
    if comment:
        run2 = p.add_run(f" {comment}")
        run2.font.size = Pt(10)
        run2.italic = True
    return p

def add_correction(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"CORRECTION: {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RED
    run.italic = True
    return p

def add_ai_flag(doc, text="AI-GENERATED - No marks awarded"):
    p = doc.add_paragraph()
    run = p.add_run(f"FLAG: {text}")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RED
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"NOTE: {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = BLUE
    run.italic = True
    return p

doc = docx.Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
pf = style.paragraph_format
pf.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EXECUTIVE FELLOWSHIP IN AI LAW & DIGITAL SOVEREIGNTY")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Module 1 - Week 3 - Databases")
run.bold = True
run.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CANDIDATE: INGRID KATABAZI")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 0, 120)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Formative Assessment - 50 Marks")
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("FINAL SCORE: 46 / 50 (92%) - GRADE A")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = GREEN

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Marking Guide Key: [Marks Awarded/Marks Available] | CORRECTION in red | FLAG in red = AI concern")
run.font.size = Pt(10)
run.italic = True

doc.add_paragraph()

# ===== QUESTION 1 =====
add_comment(doc, "QUESTION 1: Which Record to Believe? (12 marks)")
doc.add_paragraph()

add_mark(doc, "Q1(a) - ACID/BASE Explanation", 4, 4, "Excellent plain-English explanation. Clear, structured, uses concrete examples. Accounting book analogy is effective.")
add_normal(doc, "Ingrid's Answer (Q1a): Systematically breaks down ACID (Atomicity, Consistency, Isolation, Durability) and BASE (Basically Available, Soft State, Eventual Consistency). Uses analogies: 'PostgreSQL behaves like a highly controlled accounting book.' Explains replication delay between Mbarara and Kampala nodes.")
add_correction(doc, "No technical errors. This is a model answer.")
doc.add_paragraph()

add_mark(doc, "Q1(b) - Reliability Analysis", 3.5, 4, "Correct choice (PostgreSQL). Strong technical analysis. Weakness: cites ETA s.4 and s.5 instead of the directly relevant s.8.")
add_normal(doc, "Ingrid's Answer (Q1b): 'PostgreSQL record is the more reliable primary evidence. It is the system of record for financial transactions. ACID compliance provides stronger transaction integrity. Under section 4 of the ETA... section 5... The court will consider how the record was generated, whether the system operated reliably.' Also applies all 10 Amongin steps.")
add_correction(doc, "Good analysis, but the ETA section references need correction: you cite s.4 (legal recognition of data messages) and s.5 (writing requirement) when the directly relevant provision is s.8 (admissibility and evidential weight of data messages). Section 8(3) deals with the best evidence rule and s.8(4) with weight. Section 8(5) with the presumption of proper operation. These are the core provisions for this question.")
add_note(doc, "The Amongin 10-step application at the end of Q1(b) is comprehensive but appears mechanically inserted. The steps are correctly applied but the structure mirrors AI-generated formatting patterns. Student input is evident in the content (errors in ETA section numbers suggest genuine understanding gaps). Awarding marks.")
doc.add_paragraph()

add_mark(doc, "Q1(c) - BigQuery Analysis", 4, 4, "Excellent. Identifies 4 verification areas (source data, transformation, integrity, audit trail). Practical, specific, well-organized.")
add_normal(doc, "Ingrid's Answer (Q1c): Identifies ETL process (Extract, Transform, Load). Lists 4 verification areas: source data verification, transformation process scrutiny, warehouse record integrity, audit trail preservation. Recommends ETL logs, extraction timestamps, transformation rules, database synchronisation records.")
add_correction(doc, "Excellent. Practical and specific. No corrections needed.")
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Q1 Total: 11.5 / 12")
run.bold = True
doc.add_paragraph()

# ===== QUESTION 2 =====
add_comment(doc, "QUESTION 2: Admissibility Challenge (14 marks)")
doc.add_paragraph()

add_mark(doc, "Q2(a) - Objection 1 (Originality)", 4, 5, "Proper court submission format. References s.6 and s.8. Good explanation of ACID guarantees. Treats s.6 as addressing originality (following question's lead) - this is incorrect but the question itself misdirects.")
add_normal(doc, "Ingrid's Answer (Q2a): Sets out formal written submissions. 'Section 6... provides that where the law requires information to be presented or retained in its original form... Section 8 recognizes the legal validity of retaining information electronically.' References ACID guarantees and records officer certification.")
add_correction(doc, "The answer follows the question's instruction to 'consider sections 6 and 8' but note: Section 6 of the ETA deals with electronic signatures, NOT originality. The originality provision is Section 7 (Authenticity of a data message). The question itself contained this error. In practice, always cite s.7 for originality/authenticity arguments. Also, 'section 8 recognizes the legal validity of retaining information electronically' - this is actually s.9 (Retention), not s.8.")
doc.add_paragraph()

add_mark(doc, "Q2(b) - Objection 2 (Hearsay)", 4, 5, "Good hearsay analysis. Correctly distinguishes machine-generated from human statements. References s.46 Evidence Act - creative but slightly off. Misses Spiby case name.")
add_normal(doc, "Ingrid's Answer (Q2b): 'The Cassandra printout is not a narrative or statement created by the data analyst. It is a machine-generated electronic record. Under section 46 of the Evidence Act, opinions of persons skilled in science or art are admissible... The analyst functions like a records officer producing a certified extract from a register.'")
add_correction(doc, "Good reasoning. Two refinements: (1) Cite R v Spiby (1990) by name - this is the leading UK authority distinguishing real evidence from hearsay for computer output. (2) Section 46 Evidence Act deals with expert opinion on foreign law - the correct section for experts in science/art is s.43 (Opinions of experts). However, the key point (machine records are not hearsay) is correctly argued.")
doc.add_paragraph()

add_mark(doc, "Q2(c) - Objection 3 (ETL Pipeline)", 4, 4, "Excellent. Names Dremel and Colossus specifically. Lists comprehensive audit log types. Acknowledges legitimate concern. Clear structure.")
add_normal(doc, "Ingrid's Answer (Q2c): Lists: ETL execution logs, source table references, timestamps, job execution history, user identities, Dremel query execution logs, BigQuery access/modification logs. References Colossus distributed storage and immutability.")
add_correction(doc, "This is a model answer for this section. Comprehensive, specific, and practical. No corrections.")
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Q2 Total: 12 / 14")
run.bold = True
doc.add_paragraph()

# ===== QUESTION 3 =====
add_comment(doc, "QUESTION 3: Cross Examination (12 marks)")
doc.add_paragraph()

add_mark(doc, "Q3(a) - Strategy (5-Step Framework)", 3, 3, "Excellent. Explicitly follows the 5-step framework with each step labeled and applied. This is exactly what the question required.")
add_normal(doc, "Ingrid's Answer (Q3a): 'Step 1: Identify the Legal and Factual Issue - whether the Cassndra printout is reliable... Step 2: Identify Applicable Law and Technical Principles - ETA sections 6 and 8, BASE model, eventual consistency... Step 3: Apply the Law and Technical Facts... Step 4: Evaluate Strengths and Weaknesses... Step 5: Reach a Practical Litigation Strategy.'")
add_correction(doc, "Excellent use of the 5-step framework. No corrections needed. Note: typo 'Cassndra' - minor.")
doc.add_paragraph()

add_mark(doc, "Q3(b) - Four Cross-Examination Questions", 4, 4, "Excellent. Four well-targeted questions, each with a different focus (replication delay, single node, timing, verification). Each has a clear stated purpose.")
add_normal(doc, "Ingrid's Questions: Q1 'Cassandra operates on an eventual consistency model, meaning updates made to one node are not reflected immediately on all other nodes?' - Purpose: establish possibility of inconsistent records. Q2 'The printout was generated by querying only the Kampala node?' - Purpose: establish printout represents only one node's view. Q3 'The printout was not produced on 2 June but on 10 June?' - Purpose: establish not contemporaneous. Q4 'Did you verify replication status and synchronisation logs?' - Purpose: establish no verification.")
add_correction(doc, "All four questions are well-constructed, non-leading (appropriately open), and target different dimensions of unreliability. Model answer quality.")
doc.add_paragraph()

add_mark(doc, "Q3(c) - Amongin Steps", 2.5, 3, "Identifies Steps 4, 5, 6. Applies them correctly. Only 3 of the most relevant steps identified - could include Step 10 as well.")
add_normal(doc, "Ingrid's Amongin: 'Step 4 - Reliability of the Electronic System: Ask whether Cassandras BASE model and eventual consistency allow different nodes to show different information... Step 5 - Integrity of the Electronic Record: Ask whether the printout produced on 10 June accurately represented the database state on 2 June... Step 6 - Audit Trails and Metadata: Ask whether SendIt produced replication logs and synchronisation records.'")
add_correction(doc, "Good identification of relevant steps. Could also include Step 10 (third-party reproducibility) - this directly applies to the replication lag issue: an independent party querying all three nodes would have obtained different results, proving the printout is unreliable.")
doc.add_paragraph()

add_mark(doc, "Q3(d) - Answer to Judge on Weight", 2, 2, "Excellent. References s.8(4) correctly. Clear and persuasive.")
add_normal(doc, "Ingrid's Answer: 'Section 8(4) of the ETA requires the Court to consider the reliability of the manner in which the record was generated, stored and communicated. Record 2 is admissible but its weight should be reduced because it came from a BASE consistency Cassandra database where different nodes may temporarily contain different information.'")
add_correction(doc, "Perfect. No corrections.")
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Q3 Total: 11.5 / 12")
run.bold = True
doc.add_paragraph()

# ===== QUESTION 4 =====
add_comment(doc, "QUESTION 4: Data Retention and Deletion (12 marks)")
doc.add_paragraph()

add_mark(doc, "Q4(a) - Deletion Deprivation", 4, 4, "Excellent. Correctly identifies that Cassandra stores different data. Lists specific categories of lost evidence. Well-reasoned.")
add_normal(doc, "Ingrid's Answer (Q4a): 'Yes, deleting the Cassandra records could have deprived the agent of useful evidence. Cassandra and PostgreSQL do different jobs. PostgreSQL is like the official payment book; Cassandra is like notebooks in different branches where updates may reach each notebook at different times. The Cassandra records could have shown: what happened in each branch, whether there were delays in updating, whether reversals happened before the payment was recorded.'")
add_correction(doc, "Excellent. Directly addresses the architectural difference and identifies specific categories of lost evidence. No corrections.")
doc.add_paragraph()

add_mark(doc, "Q4(b) - 60-Day Retention Compliance", 3.5, 4, "Strong. References s.9 ETA, addresses Limitation Act, correctly argues each database independently subject to retention. Minor: cites Cap. 90 instead of Cap. 80.")
add_normal(doc, "Ingrid's Answer (Q4b): 'Under section 9 ETA, electronic records must be retained to remain accessible and capable of being reproduced. The Limitation Act Cap. 90 provides that contractual claims may generally be brought within 6 years. Each database must be assessed independently. PostgreSQL may preserve official payment records, but Cassandra may contain different information such as node updates and replication history.'")
add_correction(doc, "Correct analysis. Note: the Limitation Act is Cap. 80, not Cap. 90. Minor citation error but the legal principle (6-year limitation period) is correctly stated. The key argument - each database independently subject to retention - is well made.")
doc.add_paragraph()

add_mark(doc, "Q4(c) - Retention Policy Draft", 3.5, 4, "Good. PostgreSQL 7 years, Cassandra at least 6 years, BigQuery 7 years with review. References DPA s.14. Could explain why periods differ in more depth.")
add_normal(doc, "Ingrid's Answer (Q4c): 'PostgreSQL (transaction records): 7 years - supported by ETA, Limitation Act, regulatory requirements. Cassandra (operational/distributed records): at least 6 years for transaction-related data - comply with s.9 ETA. BigQuery (reports/analytics): 7 years with regular review - comply with DPA s.14. Different retention periods are justified because each database serves a different purpose.'")
add_correction(doc, "Good practical policy. Could strengthen: (1) Reference the NPS Act requirement of 10 years for payment service providers - this overrides the lower recommendations; (2) Explain more specifically why Cassandra's BASE architecture might require different treatment (replication logs may be needed to verify data state during disputes).")
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Q4 Total: 11 / 12")
run.bold = True
doc.add_paragraph()

# ===== OVERALL =====
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("=" * 50)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ASSESSMENT SUMMARY")
run.bold = True
run.font.size = Pt(13)

summary_items = [
    ("Question 1: Which Record to Believe?", "11.5 / 12", "96%"),
    ("Question 2: Admissibility Challenge", "12 / 14", "86%"),
    ("Question 3: Cross Examination", "11.5 / 12", "96%"),
    ("Question 4: Data Retention", "11 / 12", "92%"),
    ("TOTAL", "46 / 50", "92% - A"),
]

for item, score, pct in summary_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{item}: ")
    run.bold = True
    run = p.add_run(f"{score} ({pct})")

doc.add_paragraph()
add_comment(doc, "KEY STRENGTHS:")
for s in ["Exceptional technical grasp of ACID/BASE and distributed database concepts", "Well-structured, professionally formatted answers throughout", "Excellent cross-examination questions with strategic purposes", "Strong on BigQuery architecture (Dremel, Colossus, audit logs)", "Practical, actionable recommendations", "Explicitly follows frameworks as instructed (5-step, Amongin)"]:
    add_normal(doc, f"  - {s}")

doc.add_paragraph()
add_comment(doc, "KEY WEAKNESSES:")
for s in ["ETA section references occasionally imprecise (s.4/s.5 instead of s.8; Cap. 90 instead of Cap. 80)", "Followed question's misdirection on s.6 (originality) without noting it is actually about signatures", "Did not cite Spiby case name in hearsay analysis", "Section 46 Evidence Act cited instead of s.43 for expert opinion"]:
    add_normal(doc, f"  - {s}")

doc.add_paragraph()
add_comment(doc, "AI GENERATION ASSESSMENT:")
add_normal(doc, "  - Several passages show AI-assisted structure: mechanical Amongin 10-step application in Q1(b), formal court submission format in Q2(a), explicit 5-step framework labelling in Q3(a). These sections are professionally structured but the formatting and terminology closely mirror AI-generated legal writing patterns.")
add_normal(doc, "  - HOWEVER: There is clear student input throughout. The ETA section errors (s.4/s.5 instead of s.8), the Cap. 90 instead of Cap. 80 error, and the s.46 instead of s.43 error are genuine student mistakes that an AI would be unlikely to make. These indicate the student understood the material and wrote the content with AI assistance in structure/organisation only.")
add_normal(doc, "  - DECISION: Marks are AWARDED in full. The AI assistance appears limited to formatting and organisation; the substantive legal reasoning, technical knowledge, and errors are the student's own work. This is consistent with appropriate use of AI as a drafting tool (not generating substantive content).")
add_normal(doc, "  - Recommendation: Advise the student to verify ETA section numbers (s.7 for originality, s.8 for admissibility, s.9 for retention) and Evidence Act sections (s.43 for experts) before final submission. These citation errors are easily avoidable.")

path = r'C:\Users\DELL\research\Week_3_Ingrid_Marked.docx'
doc.save(path)
print(f"Ingrid marked assessment saved to {path}")

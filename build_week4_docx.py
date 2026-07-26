import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Page Setup
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Style Definitions
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level, (size, color_hex) in enumerate([
    (26, '1a3c5e'),
    (16, '2c5f8a'),
    (13, '3a7ab5'),
    (12, '4a8fc4'),
], start=1):
    s = doc.styles[f'Heading {level}']
    s.font.name = 'Calibri'
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(*bytes.fromhex(color_hex))
    s.paragraph_format.space_before = Pt(18 if level == 1 else 14)
    s.paragraph_format.space_after = Pt(8)

def add_para(text, bold=False, italic=False, size=None, color=None, align=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size: run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    if align: p.alignment = align
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def add_body(text):
    return add_para(text)

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.size = Pt(11)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.size = Pt(10)
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for para in row_cells[i].paragraphs:
                for r in para.runs:
                    r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('_' * 80)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run.font.size = Pt(8)

def add_spacer():
    doc.add_paragraph()

# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

add_para('MODULE 1: DIGITAL TECHNOLOGY FUNDAMENTALS', bold=True, size=16, color='1a3c5e', align=WD_ALIGN_PARAGRAPH.CENTER)
add_spacer()
add_para('WEEK 4', bold=True, size=22, color='1a3c5e', align=WD_ALIGN_PARAGRAPH.CENTER)
add_spacer()
add_para('APIs, Cloud Computing & SDLC', bold=True, size=28, color='2c5f8a', align=WD_ALIGN_PARAGRAPH.CENTER)
add_spacer()
add_para('Reading Notes for Legal Practitioners', italic=True, size=14, color='666666', align=WD_ALIGN_PARAGRAPH.CENTER)
add_spacer()
add_para('Executive Fellowship in AI Law & Digital Sovereignty', size=12, color='888888', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Book Manuscript \u2014 Chapter 4', size=11, color='aaaaaa', align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

# ============================================================
# CONTENTS
# ============================================================
doc.add_heading('Contents', level=1)

toc = [
    ('How to Use These Notes', True),
    ('Part 1: Application Programming Interfaces (APIs)', True),
    ('  1.1  What is an API? The Waiter Analogy', False),
    ('  1.2  RESTful Architecture', False),
    ('  1.3  API Endpoints and HTTP Verbs', False),
    ('  1.4  Payloads \u2014 JSON and XML', False),
    ('  1.5  API Authentication \u2014 OAuth, API Keys, and Tokens', False),
    ('  1.6  API Security \u2014 The OWASP API Security Top 10', False),
    ('  1.7  Real-World Example: UGHub \u2014 Uganda\'s National API Gateway', False),
    ('  1.8  Legal Framework for APIs', False),
    ('  1.9  Cross-Examination Questions for API Evidence', False),
    ('Part 2: Cloud Computing', True),
    ('  2.1  What is Cloud Computing? The Renting Analogy', False),
    ('  2.2  The Three Cloud Service Models', False),
    ('  2.3  Cloud Deployment Models', False),
    ('  2.4  Service Level Agreements (SLAs)', False),
    ('  2.5  Legal Framework for Cloud Computing', False),
    ('  2.6  Data Localisation and Cross-Border Cloud Storage', False),
    ('  2.7  Cross-Examination Questions for Cloud Disputes', False),
    ('Part 3: The Software Development Lifecycle (SDLC) & LLMOps', True),
    ('  3.1  What is the SDLC? The House-Building Analogy', False),
    ('  3.2  The Six Phases of the SDLC', False),
    ('  3.3  Waterfall vs Agile vs DevOps', False),
    ('  3.4  LLMOps \u2014 The SDLC for AI Models', False),
    ('  3.5  Legal Framework for the SDLC', False),
    ('  3.6  Cross-Examination Questions for SDLC Failures', False),
    ('Part 4: Foundation-to-Tune Comparative Analysis', True),
    ('Part 5: Comparative International Legal Frameworks', True),
    ('Part 6: Legal Problem-Solving Framework', True),
    ('  6.1  How to Analyse an API/Cloud/SDLC Problem', False),
    ('  6.2  The Week 4 Practice Task \u2014 SLA Compliance Audit', False),
    ('Part 7: Quick Reference Cards', True),
    ('Part 8: Glossary for Lawyers', True),
]

for item, is_bold in toc:
    p = doc.add_paragraph()
    run = p.add_run(item)
    run.font.size = Pt(10)
    run.bold = is_bold

doc.add_page_break()

# ============================================================
# HOW TO USE
# ============================================================
doc.add_heading('How to Use These Notes', level=1)
add_body('Each section follows the same structure:')
for s in ['Plain-English Explanation \u2014 what it is, in everyday language',
          'Practical Illustration \u2014 a concrete example tied to legal practice',
          'Why It Matters for a Lawyer \u2014 the legal significance',
          'Legal Framework \u2014 relevant Ugandan statutes, regulations, and international instruments']:
    add_bullet(s)

add_divider()

# ============================================================
# PART 1: APIs
# ============================================================
doc.add_heading('PART 1: APPLICATION PROGRAMMING INTERFACES (APIs)', level=1)

doc.add_heading('1.1  What is an API? The Waiter Analogy', level=2)
add_body('Imagine you are at a restaurant. You are sitting at your table (the client). The kitchen (the server) has all the food and capabilities. But you cannot walk into the kitchen and cook for yourself \u2014 that would be chaos. Instead, a waiter brings your order to the kitchen and returns with your food.')
add_body('The waiter is the API. It is a messenger that takes requests, tells the kitchen what to do, and returns the response to you.')

add_table(
    ['Real World', 'Technical World'],
    [['You (customer)', 'Client application (mobile app, website)'],
     ['Waiter', 'API (Application Programming Interface)'],
     ['Kitchen', 'Server (back-end system, database)'],
     ['Menu', 'API documentation (endpoints available)'],
     ['Your order', 'API request (HTTP method + endpoint + data)'],
     ['Food returned', 'API response (data in JSON/XML format)'],
     ['Kitchen rules', 'API schema and validation rules']],
    col_widths=[6, 10]
)
add_spacer()

add_body('Plain English: An API is a set of rules that allows one piece of software to talk to another. It defines what requests can be made, how to make them, and what format the response will take. It is the contract between two systems.')
add_body('Practical Illustration \u2014 MTN Mobile Money API: When a fintech app wants to check a customer\'s MTN Mobile Money balance, it sends an API request. The API verifies authorisation, checks the balance, and responds. The fintech never sees MTN\'s database.')
add_body('Why this matters for a lawyer:')
for s in ['If a transaction fails, the API log is the evidence trail',
          'If unauthorised data is exposed, the API security configuration determines liability',
          'If an API goes down, the SLA determines whether there is a remedy',
          'If an API returns incorrect data, you need to know whether the error is in the API or the system behind it']:
    add_bullet(s)

doc.add_heading('1.2  RESTful Architecture \u2014 The Most Common API Type', level=2)
add_body('REST (Representational State Transfer) is the most common API type. It is stateless (each request is self-contained), resource-based (everything has a unique URL), and uses standard HTTP methods.')
add_body('Why this matters: The stateless nature means each API call is self-contained. If you subpoena API logs, each entry contains everything needed to understand what happened.')

doc.add_heading('1.3  API Endpoints and HTTP Verbs', level=2)
add_body('An endpoint is the specific URL where an API can be accessed. Think of it as a specific desk in a government office.')
add_table(
    ['Endpoint', 'What It Does'],
    [['https://api.bank.co.ug/customers', 'Access customer records'],
     ['https://api.bank.co.ug/customers/123', 'Access a specific customer'],
     ['https://api.bank.co.ug/transactions', 'Access transaction records'],
     ['https://api.bank.co.ug/payments', 'Initiate or view payments']],
    col_widths=[8, 8]
)
add_spacer()

add_body('HTTP Verbs tell the API what action to perform:')
add_table(
    ['Verb', 'Action', 'Legal Analogy'],
    [['GET', 'Retrieve data', 'Requesting a court file'],
     ['POST', 'Create new data', 'Filing a new claim'],
     ['PUT', 'Replace existing data', 'Amending entire pleadings'],
     ['PATCH', 'Partially update data', 'Correcting a typo'],
     ['DELETE', 'Remove data', 'Withdrawing a claim']],
    col_widths=[3, 5, 8]
)
add_spacer()
add_body('Forensic significance: A GET to /customers with no filters = data scraping. A DELETE to /transactions = evidence tampering.')

doc.add_heading('1.4  Payloads \u2014 JSON and XML', level=2)
add_body('The payload is the actual data carried by the API message. JSON is the modern standard \u2014 lightweight, key-value pairs. XML is older, more verbose, still used in banking and government.')
add_body('Forensic question: Was the data in the API payload exactly what the user authorised, or could additional fields have been included without consent? This is a data minimisation issue under the DPA.')

doc.add_heading('1.5  API Authentication \u2014 OAuth, API Keys, and Tokens', level=2)
add_bullet('API Keys: Simple identifier string. Like a building access card \u2014 anyone holding it can enter.')
add_bullet('OAuth 2.0: Temporary tokens. Like a visitor\'s pass that expires at 5:00 PM.')
add_bullet('JWT: Self-contained, digitally signed token. Like a notarised letter of introduction.')
add_body('Why it matters: If a data breach occurs via API, the first question is how authentication worked. Under DPA Section 20, authentication method choice is directly relevant to whether the data controller discharged their duty.')

doc.add_heading('1.6  API Security \u2014 The OWASP API Security Top 10', level=2)
add_body('The OWASP API Security Top 10 (2023) is the industry standard:')
add_table(
    ['Rank', 'Risk', 'What It Means'],
    [['API1', 'Broken Object Level Authorization', 'User A can access User B\'s data'],
     ['API2', 'Broken Authentication', 'Weak credentials allow unauthorised access'],
     ['API3', 'Broken Object Property Level Authorization', 'Excessive data exposure'],
     ['API4', 'Unrestricted Resource Consumption', 'No rate limiting'],
     ['API5', 'Broken Function Level Authorization', 'Regular user accesses admin functions'],
     ['API6', 'Unrestricted Access to Sensitive Flows', 'Bots abuse business logic'],
     ['API7', 'Server Side Request Forgery', 'Attacker tricks server into internal requests'],
     ['API8', 'Security Misconfiguration', 'Default passwords, unpatched systems'],
     ['API9', 'Improper Inventory Management', 'Deprecated APIs still accessible'],
     ['API10', 'Unsafe Consumption of APIs', 'Client blindly trusts third-party APIs']],
    col_widths=[2, 5, 9]
)
add_spacer()
add_body('Why it matters: OWASP Top 10 constitutes "generally accepted information security practices" under DPA Section 20(3). Failure to address these risks is evidence of non-compliance.')

doc.add_heading('1.7  Real-World Example: UGHub \u2014 Uganda\'s National API Gateway', level=2)
add_body('UGHub is the Government of Uganda\'s enterprise API gateway, operated by NITA-U, running on WSO2 API Manager. As of 2026, it connects over 135 government entities, providing secure data exchange between MDAs.')
add_body('Why it matters: If litigating against a government agency, UGHub logs may contain the evidence. If advising a contractor integrating with government data, they must comply with UGHub\'s API security requirements.')

doc.add_heading('1.8  Legal Framework for APIs', level=2)

doc.add_heading('Electronic Transactions Act, Cap. 99, Sections 29\u201333', level=3)
add_bullet('Section 29: Service provider not liable for third-party material if merely providing access. Exception: contractual obligations remove the exemption.')
add_bullet('Section 30: No liability for linking to infringing material if no actual knowledge and access removed after notification.')
add_bullet('Section 31: Notice-and-takedown procedure with formal notification requirements.')
add_bullet('Section 32: No duty to monitor transmitted data.')
add_body('Practical tip: An API provider that merely transmits data may benefit from the S.29 exemption \u2014 but not if a contract exists. This is why cloud SLAs matter: they create contractual obligations that remove the exemption.')

doc.add_heading('Data Protection and Privacy Act, 2019, Sections 20\u201322', level=3)
add_bullet('Section 20: Mandates appropriate, reasonable, technical and organisational security measures. Requires risk identification, safeguard implementation, regular verification, and continuous updates.')
add_bullet('Section 21: Controller-processor contract required. The cloud DPA is a statutory requirement, not optional.')
add_bullet('Section 22: Operator/authorised person confidentiality obligations.')

doc.add_heading('Data Protection and Privacy Regulations, 2021, Regulations 31\u201333', level=3)
add_bullet('Regulation 31: Publish security practices and procedures.')
add_bullet('Regulation 32: Specific measures \u2014 access control, encryption, system monitoring.')
add_bullet('Regulation 33: Mandatory breach notification to PDPO.')

doc.add_heading('Uganda Communications Act, Cap. 103', level=3)
add_body('Conditional application: UCC Act applies to licensed communications operators. If the API provider is a licensed operator, UCC (Interconnection and Access) Regulations 2019 and UCC (Quality of Service) Regulations 2019 add additional obligations. Does NOT apply to standard fintech/healthtech using APIs for internal operations.')

doc.add_heading('1.9  Cross-Examination Questions for API Evidence', level=2)
for cat, qs in [
    ('Authentication & Access', ['"What authentication method did your API use?"', '"Can you produce the API access logs for the relevant period?"']),
    ('API Security', ['"Had you conducted an API security assessment based on OWASP Top 10?"', '"Was your API tested for Broken Object Level Authorization?"']),
    ('ETA Section 29 Liability', ['"Was the API merely transmitting data, or processing/storing data?"', '"Did you have a contractual obligation beyond mere transmission?"']),
]:
    doc.add_heading(cat, level=3)
    for q in qs:
        add_bullet(q)

add_divider()

# ============================================================
# PART 2: CLOUD COMPUTING
# ============================================================
doc.add_heading('PART 2: CLOUD COMPUTING', level=1)

doc.add_heading('2.1  What is Cloud Computing? The Renting Analogy', level=2)
add_table(
    ['Model', 'Analogy', 'You Manage', 'Provider Manages'],
    [['On-Premise', 'Own your building', 'Everything', 'Nothing'],
     ['IaaS', 'Rent empty office', 'Apps, data, OS, middleware', 'Servers, storage, networking'],
     ['PaaS', 'Rent furnished office', 'Applications, data', 'Everything else'],
     ['SaaS', 'Rent serviced office', 'Nothing', 'Everything']],
    col_widths=[3, 5, 5, 5]
)
add_spacer()
add_body('Plain English: Cloud computing delivers computing services over the internet. You pay only for what you use, like a utility bill.')

doc.add_heading('2.2  The Three Cloud Service Models', level=2)
add_body('IaaS: Rent raw computing resources. Most customer control, most responsibility. DPA S.20 compliance falls primarily on the customer.')
add_body('PaaS: Rent a complete platform. Shared liability. DPA S.21 controller-processor contract is critical.')
add_body('SaaS: Use a complete application. Provider has most control and responsibility. Must include a Data Processing Agreement.')

doc.add_heading('2.3  Cloud Deployment Models', level=2)
add_bullet('Public Cloud: Shared infrastructure. Legal concern: data on shared infrastructure with other tenants.')
add_bullet('Private Cloud: Dedicated infrastructure. Full control over data location and security.')
add_bullet('Sovereign Cloud: Physically located within a country\'s borders, subject exclusively to that country\'s laws.')
add_body('East African sovereign cloud initiatives: Servernah (Nairobi), Savannah Cloud (Nairobi), Konza National Data Centre (Kenya), Karuma AI Supercomputing Hub (Uganda), EAC Cloud (Arusha), UniCloud Africa (pan-African).')

doc.add_heading('2.4  Service Level Agreements (SLAs)', level=2)
add_body('An SLA is the contractual promise of service quality \u2014 measurable targets with remedies for failure.')
add_table(
    ['Metric', 'What It Measures', 'Typical Promise'],
    [['Uptime', 'Service accessibility', '99.9% (~8.7 hrs/year downtime)'],
     ['Latency', 'Response speed', '95th percentile < 200ms'],
     ['Throughput', 'Request capacity', '10,000 requests/min'],
     ['Error Rate', 'Failed requests', '< 0.1%'],
     ['Recovery Time', 'Restore speed', '< 4 hours']],
    col_widths=[4, 5, 7]
)
add_spacer()
add_body('The SLA credit trap: Most cloud SLAs give service credits (percentage of monthly fee), not cash damages. If UGX 50M in losses result from an outage, the SLA credit might be UGX 5M \u2014 and it is the exclusive remedy.')
add_body('Cross-examination: "Is the SLA credit the sole and exclusive remedy, regardless of actual loss?"')

doc.add_heading('2.5  Legal Framework for Cloud Computing', level=2)
add_bullet('DPA Section 21: Every cloud agreement involving personal data must include a DPA specifying security measures, sub-processing restrictions, breach notification, and data return/deletion.')
add_bullet('ETA Sections 29\u201333: The "mere conduit" exemption is lost where a contractual obligation exists.')
add_bullet('UCC Quality of Service Regulations 2019: For licensed operators, minimum service standards override weaker contractual SLA terms.')
add_bullet('UNCITRAL Notes on Cloud Computing Contracts (2019): International best practice on liability allocation, measurable SLAs, audit rights, and termination/transition provisions.')

doc.add_heading('2.6  Data Localisation and Cross-Border Cloud Storage', level=2)
add_table(
    ['Jurisdiction', 'Data Localisation', 'Source'],
    [     ['Uganda', 'No explicit localisation law. Cross-border transfers restricted without adequacy or consent.', 'DPA Section 19'],
     ['Kenya', 'Localisation for certain categories. Cloud Policy 2025 in force.', 'Kenya Cloud Policy 2025; DPA S.50'],
     ['Rwanda', 'Enforced. MTN fined USD 8.2M for transferring data outside Rwanda.', 'RURA Regulations 2016'],
     ['EAC', 'Harmonised framework validated June 2026 (not yet adopted).', 'EAC/EARDIP'],
     ['AfCFTA', 'Art. 22 prohibits mandatory localisation (with exceptions). Not ratified.', 'Digital Trade Protocol 2024']],
    col_widths=[3, 8, 5]
)

doc.add_heading('2.7  Cross-Examination Questions for Cloud Disputes', level=2)
for cat, qs in [
    ('SLA Failures', ['"What was the promised uptime percentage?"', '"Can you produce the uptime calculation for the month in dispute?"']),
    ('Data Location', ['"In which country was the data physically stored?"', '"Did the provider have the right to transfer data to sub-processors elsewhere?"']),
    ('Security', ['"What certifications does the provider hold (ISO 27001, SOC 2)?"', '"Can you produce the security audit report?"']),
    ('Liability', ['"Is the SLA credit the exclusive remedy?"', '"Does the limitation clause exclude data protection claims?"']),
]:
    doc.add_heading(cat, level=3)
    for q in qs:
        add_bullet(q)

add_divider()

# ============================================================
# PART 3: SDLC & LLMOps
# ============================================================
doc.add_heading('PART 3: THE SOFTWARE DEVELOPMENT LIFECYCLE (SDLC) & LLMOps', level=1)

doc.add_heading('3.1  What is the SDLC? The House-Building Analogy', level=2)
add_body('Just as a house is built in stages, software is built in phases. Each phase generates evidence of what was done, by whom, and whether proper procedures were followed.')
add_table(
    ['Phase', 'House Analogy', 'What Happens'],
    [['Requirements', 'Architect meets client', 'Determine what software must do'],
     ['Design', 'Draw blueprints', 'Plan software structure'],
     ['Development', 'Build', 'Write code'],
     ['Testing/QA', 'Inspector checks', 'Verify software works'],
     ['Deployment', 'Client moves in', 'Release to users'],
     ['Maintenance', 'Repairs', 'Fix bugs, patches, new features']],
    col_widths=[4, 5, 7]
)

doc.add_heading('3.2  The Six Phases of the SDLC', level=2)
phases = [
    ('Requirements', 'Determine what the software must do. Produces a Requirements Specification.',
     'Baseline for determining if software performed as intended. Privacy requirements should be specified here.',
     '"Does the requirements specification mention data protection? If not, why not?"'),
    ('Design', 'Create architecture, database schema, API design, security architecture.',
     'Privacy-by-design and security-by-design must be embedded here. DPA S.20(2)(a) requires risk identification.',
     '"Was a DPIA conducted during the design phase? Can you produce it?"'),
    ('Development', 'Write the code.',
     'IP ownership of code. Open source licence compliance. Liability for AI-assisted coding.',
     '"Was third-party or open-source code used? What licence terms applied?"'),
    ('Testing / QA', 'Unit, integration, security, and user acceptance testing.',
     'DPA S.20(2)(c) requires regular verification. NIST SP 800-228 recommends API security testing throughout.',
     '"Can you produce test reports? Was security testing conducted? What vulnerabilities were found?"'),
    ('Deployment', 'Release to users.',
     'Under ETA S.8, deployment logs are admissible evidence showing which version was running.',
     '"Can you produce the deployment log for the disputed period?"'),
    ('Maintenance', 'Fix bugs, apply patches, add features.',
     'DPA S.20(2)(d) requires continually updating safeguards. Failure to patch known vulnerabilities increases liability.',
     '"When was the vulnerability first discovered? When was a patch available? When was it applied?"'),
]
for title, what, legal, crossq in phases:
    doc.add_heading(f'Phase: {title}', level=3)
    add_body(f'What happens: {what}')
    add_body(f'Legal significance: {legal}')
    add_body(f'Cross-examination: {crossq}')

doc.add_heading('3.3  Waterfall vs Agile vs DevOps', level=2)
add_table(
    ['Aspect', 'Waterfall', 'Agile', 'DevOps'],
    [['Planning', 'Complete upfront', 'Iterative', 'Continuous'],
     ['Documentation', 'Comprehensive', 'Minimal', 'Automated'],
     ['Testing', 'Dedicated phase', 'Throughout sprints', 'Continuous/automated'],
     ['Deployment', 'Single event', 'Every sprint', 'Continuous'],
     ['Change cost', 'High', 'Low', 'Very low'],
     ['Best for', 'Fixed requirements', 'Evolving products', 'Cloud-native'],
     ['Legal risk', 'Scope disputes', 'Scope creep', 'Deployment control']],
    col_widths=[3, 4, 4, 4]
)

doc.add_heading('3.4  LLMOps \u2014 The SDLC for AI Models', level=2)
add_body('LLMOps adapts the SDLC for AI models with additional governance layers: data lineage, bias detection, explainability, human-in-the-loop, hallucination detection, and prompt security.')
add_body('Why it matters: If an AI model deployed via API gives incorrect legal advice, the LLMOps governance log determines liability. Kenya AI Bill 2026, Section 26 requires data logging for high-risk AI systems.')

doc.add_heading('3.5  Legal Framework for the SDLC', level=2)
add_body('DPA Section 20 transforms the SDLC from a technical choice into a legal obligation:')
add_table(
    ['Phase', 'DPA Requirement', 'In Practice'],
    [['Requirements', 'Identify risks S.20(2)(a)', 'Document privacy risks upfront'],
     ['Design', 'Establish safeguards S.20(2)(b)', 'Security-by-design from the start'],
     ['Testing', 'Verify effectiveness S.20(2)(c)', 'Documented security testing'],
     ['Deployment', 'Organisational measures S.20(1)', 'Security checks in deployment'],
     ['Maintenance', 'Update safeguards S.20(2)(d)', 'Timely patch application']],
    col_widths=[4, 5, 7]
)
add_spacer()
add_body('NITA-U NISF 2026: Launched July 2026. Minimum baseline security controls for critical infrastructure. Government contractors must comply. References ISO 27001 Annex A.8.25 (Secure Development Lifecycle).')
add_body('ETA Section 7: Electronic records qualify as "originals" if integrity is assured. Version control (e.g., Git) provides integrity assurance for SDLC documentation.')
add_body('Evidence Act Cap. 6: SDLC documentation is a "document" under S.2, with presumptive genuineness under S.78 if certified.')

doc.add_heading('3.6  Cross-Examination Questions for SDLC Failures', level=2)
for cat, qs in [
    ('Requirements', ['"Can you produce the requirements specification for the disputed feature?"', '"Was a DPIA conducted?"']),
    ('Design', ['"Was a security architecture review conducted?"', '"Were data minimisation principles embedded?"']),
    ('Testing', ['"What testing was conducted before deployment?"', '"Was the specific vulnerability tested for?"']),
    ('Deployment', ['"Who authorised the deployment?"', '"Was a rollback plan in place?"']),
    ('Maintenance', ['"When was the vulnerability discovered?"', '"Why was the patch not applied before the breach?"']),
]:
    doc.add_heading(cat, level=3)
    for q in qs:
        add_bullet(q)

add_divider()

# ============================================================
# PART 4: FOUNDATION-TO-TUNE
# ============================================================
doc.add_heading('PART 4: FOUNDATION-TO-TUNE COMPARATIVE ANALYSIS', level=1)
add_body('The Global/Engineering Foundation: Globally, the SDLC is engineered to minimise deployment friction. API security is a risk management choice. Cloud SLAs are commercial terms.')
add_body('The Ugandan Practice Tune: Multiple statutes transform these choices into legal obligations.')
add_table(
    ['Issue', 'Global Foundation', 'Ugandan Tune'],
    [['API security', 'Risk management choice', 'DPA S.20 legal obligation'],
     ['Cloud SLA', 'Commercial term', 'DPA S.21 + UCC QoS overlay'],
     ['SDLC documentation', 'Best practice', 'Evidence Act + ETA S.7 requirement'],
     ['Data location', 'Commercial negotiation', 'Emerging localisation laws'],
     ['Provider liability', 'Contractual allocation', 'ETA S.29-33 framework'],
     ['Breach response', 'Contractual notice', 'DPA Reg 33 mandatory notification']],
    col_widths=[4, 5, 7]
)

add_divider()

# ============================================================
# PART 5: COMPARATIVE JURISDICTION TABLE
# ============================================================
doc.add_heading('PART 5: COMPARATIVE INTERNATIONAL LEGAL FRAMEWORKS', level=1)
add_body('How major jurisdictions regulate the topics covered in this week:')

add_body('Each row compares one topic across jurisdictions. Uganda is the benchmark column so you can see at a glance how other systems differ.')
add_table(
    ['Topic', 'Uganda (Benchmark)', 'United States', 'United Kingdom', 'European Union', 'South Africa', 'Kenya', 'Singapore'],
    [
        ['API / Intermediary Liability',
         'ETA S.29: no liability if merely providing access to third-party material (lost if contract exists). S.30: no liability for links. S.31: notice-and-takedown. S.32: no duty to monitor. S.29(2)(a): contractual exception.',
         'Broader immunity than Uganda. CDA S.230 treats platforms as publishers \u2014 not liable for user content at all. No equivalent to Uganda\u2019s S.29(2)(a) contractual exception. DMCA for copyright DMCA for copyright. No API-specific law.',
         'Framework similar to Uganda. ECA 2000 Reg. 17\u201319 (mere conduit, caching, hosting). No contractual exception like ETA S.29(2)(a). Defamation Act 2013 S.5 adds intermediary defence.',
         'Same roots as Uganda (both from UNCITRAL/ECD model). DSA (2022) adds obligations on large platforms Uganda lacks. No monitoring duty same as ETA S.32.',
         'Nearly identical to Uganda. ECTA S.70\u201378 mirrors ETA S.29\u201332 (same UNCITRAL model). Notice-and-takedown, no monitoring, contractual exception.',
         'No equivalent to Uganda\u2019s ETA. KICA regulates ICT generally. DPA 2019 covers data breaches via API. AI Bill S.26 adds AI logging. No safe harbour framework.',
         'No dedicated intermediary regime like Uganda. ETA covers electronic contracts. CMA criminalises unauthorised access. PDPA covers data breaches.'],
        ['Cloud / Data Localisation',
         'DPA S.19: cross-border transfers allowed only if recipient country has adequate protection or data subject consents. No explicit localisation law.',
         'No federal localisation (unlike Uganda\u2019s adequacy model). CLOUD Act gives US law enforcement access to data held by US companies anywhere. CCPA adds state-level privacy.',
         'Similar adequacy model to Uganda (UK GDPR Art. 45\u201349). UK-US Data Bridge for transfers. No localisation requirement.',
         'GDPR Ch. V adequacy/SCC/BCR same logic as Uganda S.19 but more detailed. EUCS proposed (sovereign cloud certification) \u2014 goes beyond Uganda. GAIA-X.',
         'Same approach as Uganda. POPIA S.72 requires adequate protection, consent, or binding contract for transfers. No localisation.',
         'Goes beyond Uganda. DPA S.50 requires processing through a server or data centre in Kenya (explicit localisation). Part VI restricts transfers abroad.',
         'Lighter than Uganda. PDPA requires notification and consent for transfers. IMDA Cloud Guidelines. No adequacy requirement.'],
        ['Data Protection / Security Measures',
         'DPA S.20: technical and organisational security measures (risk ID, safeguards, verification, updates). S.21: processor contract mandatory. S.22: operator confidentiality. S.23: breach notification to PDPO.',
         'Sectoral, not omnibus like Uganda. HIPAA (health), GLBA (finance), COPPA (children). FTC S.5 catches unfair practices. No single DPA equivalent. No mandatory processor contract like Uganda S.21.',
         'UK GDPR + DPA 2018 \u2014 same structure as Uganda but far more detailed. S.20\u201322 DPA 2018 parallel Uganda DPA. ICO enforcement. Higher fines (\u00a317.5M/4% turnover vs Uganda\u2019s lower penalties).',
         'GDPR \u2014 template Uganda\u2019s DPA was modelled on. Art. 24\u201343 (security, DPA, DPIA, breach notification) more comprehensive than Uganda. EDPB guidance. 4% turnover fines.',
         'POPIA S.19\u201322 virtually identical to Uganda DPA S.20\u201323 (same EU model origin). Information Regulator enforces. Similar penalty structure.',
         'DPA 2019 modelled on EU law like Uganda. S.41 adds explicit DPBD requirement (not in Uganda DPA). Part IV (S.30\u201343) similar scope. Enforcement: Data Commissioner.',
         'PDPA different structure from Uganda: 9 obligations without Uganda\u2019s explicit S.21 processor contract requirement. 2024 amendments raised fines to 10% turnover (far above Uganda).'],
        ['SDLC / Software Liability',
         'No specific SDLC law. DPA S.20 applies indirectly (security during processing). ETA S.7 (originals for version control). Evidence Act S.2, 78 (documentation).',
         'No federal SDLC statute (same gap as Uganda). FTC enforcement against deficient security practices. NIST SP 800-218 (SSDF) as guidance. Tort law evolving for software defects.',
         'Consumer Protection Act 1987 treats software as product in some contexts \u2014 beyond Uganda. ICO privacy-by-design guidance. No dedicated SDLC statute.',
         'Ahead of Uganda. Revised Product Liability Directive (2024) explicitly covers software. Cyber Resilience Act mandates secure SDLC for connected devices. NIS2 for critical sectors.',
         'Same gap as Uganda. No SDLC-specific law. Common law delict/contract. ECTA for electronic evidence.',
         'Ahead of Uganda on AI but same gap on general SDLC. DPA S.41: DPBD required. AI Bill S.26: high-risk AI obligations (risk assessments, documentation, logging).',
         'No SDLC legislation (same as Uganda). Consumer Protection (Fair Trading) Act for defective products. IMDA IoT guidelines. PDPA by-design principles.'],
        ['AI Governance',
         'No AI-specific legislation. DPA S.20 applies to AI processing of personal data. NITA-U NISF 2026 references secure development lifecycle.',
         'No comprehensive AI Act. AI Bill of Rights (non-binding). Exec. Order 14110 (testing/reporting). NIST AI RMF. State laws patchwork. Same gap as Uganda in binding law.',
         'AI Safety Institute. Pro-innovation White Paper (no binding law yet). Bletchley Declaration. Similar to Uganda \u2014 guidance without binding obligations.',
         'Far ahead. EU AI Act binding, risk-based (unacceptable, high, limited, minimal). High-risk: risk assessments, human oversight, logging, transparency. Effective 2024\u20132027 phased.',
         'Same as Uganda. No AI Act. POPIA governs AI personal data processing. Draft National AI Policy Framework.',
         'Ahead of Uganda. AI Bill 2026 risk-based (modelled on EU AI Act). S.26: risk assessments, human rights impact assessments, record-keeping, transparency for high-risk AI systems.',
         'Different approach from Uganda. Voluntary Model AI Governance Framework. AI Verify testing toolkit. Proposed regulation (2025\u20132026) \u2014 sectoral rather than omnibus.'],
    ],
    col_widths=[2.5, 4, 4, 4, 4, 4, 4, 4]
)

add_divider()

# ============================================================
# PART 6: PROBLEM-SOLVING
# ============================================================
doc.add_heading('PART 6: LEGAL PROBLEM-SOLVING FRAMEWORK', level=1)

doc.add_heading('6.1  How to Analyse an API/Cloud/SDLC Problem', level=2)
add_body('Step 1 \u2014 Identify What Failed: API failure, cloud failure, or SDLC failure?')
add_body('Step 2 \u2014 Map the Legal Framework: ETA S.29-33, DPA S.20-22, UCC QoS, UNCITRAL Notes, Evidence Act.')
add_body('Step 3 \u2014 Assess Compliance: Was an API security assessment done? Was a DPA in place? Was a DPIA conducted?')
add_body('Step 4 \u2014 Determine Liability: API provider (ETA S.29 exemption?), Cloud provider (DPA S.21 processor?), Data controller (DPA S.20 primary), SDLC team (contractual).')
add_body('Step 5 \u2014 Challenge or Authenticate Evidence: Challenge (no logs, no testing) or Authenticate (certifications, audit trails).')

doc.add_heading('6.2  The Week 4 Practice Task \u2014 SLA Compliance Audit', level=2)
add_body('Scenario: Your client, a Ugandan healthtech company, is outsourcing hosting to a multinational cloud provider (PaaS) via API gateway. The boilerplate contract has "as-is" disclaimers, "industry standard" uptime, no API metrics, SLA credits as exclusive remedy, no DPA.')
add_body('Step 1 \u2014 Legal Issues:')
add_bullet('DPA S.21 requires a controller-processor contract \u2014 missing DPA is a statutory violation')
add_bullet('DPA S.20 requires appropriate measures \u2014 "as-is" disclaimers are inconsistent')
add_bullet('ETA S.29(2)(a): contractual obligation removes liability exemption')
add_bullet('UCC QoS may apply if platform qualifies as data communication service')

add_body('Step 2 \u2014 Drafted Amendments:')
add_table(
    ['Issue', 'Problem', 'Amendment'],
    [['Uptime', '"Industry standard"', 'Define: 99.9% monthly, excluding scheduled maintenance with 48h notice'],
     ['API latency', 'Not addressed', 'Define: 95th percentile \u2264 200ms, measured monthly'],
     ['Downtime', 'Not defined', 'Define: error rate > 1% or latency > 500ms for > 5 consecutive minutes'],
     ['Remedy', 'SLA credits exclusive', 'Preserve right to actual damages for data breach, DPA non-compliance, gross negligence'],
     ['DPA', 'None attached', 'Attach DPA: data types, security measures, sub-processor restrictions, breach notification, data return'],
     ['Data location', 'Not specified', 'Primary storage in Uganda or approved jurisdiction. No transfers without consent + TIA']],
    col_widths=[3, 4, 9]
)
add_spacer()
add_body('Step 3 \u2014 Cross-Reference with ETA: Under ETA S.29(2)(a), once a contractual obligation exists, the "mere conduit" exemption is lost. The amendment should state: "By agreeing to specific technical metrics, the cloud provider accepts contractual liability. The ETA S.29 exemption does not apply because S.29(2)(a) preserves contractual obligations."')

add_divider()

# ============================================================
# PART 7: QUICK REFERENCE CARDS
# ============================================================
doc.add_heading('PART 7: QUICK REFERENCE CARDS', level=1)

doc.add_heading('HTTP Verbs', level=2)
add_table(
    ['Verb', 'Action', 'Safe?', 'Idempotent?', 'Legal Analogy'],
    [['GET', 'Read', 'Yes', 'Yes', 'Requesting a court file'],
     ['POST', 'Create', 'No', 'No', 'Filing a new claim'],
     ['PUT', 'Replace', 'No', 'Yes', 'Amending entire pleadings'],
     ['PATCH', 'Partial update', 'No', 'No', 'Correcting a typo'],
     ['DELETE', 'Remove', 'No', 'Yes', 'Withdrawing a claim']],
    col_widths=[2, 3, 2, 2, 5]
)

doc.add_heading('Cloud Service Models', level=2)
add_table(
    ['Model', 'You Manage', 'Provider Manages', 'Analogy'],
    [['On-Premise', 'Everything', 'Nothing', 'Own your building'],
     ['IaaS', 'Apps, data, OS, middleware', 'Servers, storage, networking', 'Rent empty office'],
     ['PaaS', 'Applications, data', 'Everything else', 'Rent furnished office'],
     ['SaaS', 'Nothing', 'Everything', 'Rent serviced office']],
    col_widths=[3, 5, 5, 5]
)

doc.add_heading('SLA Uptime "Nines"', level=2)
add_table(
    ['Uptime %', 'Downtime/Year', 'Downtime/Month'],
    [['99%', '3.65 days', '7.2 hours'],
     ['99.9%', '8.76 hours', '43.2 minutes'],
     ['99.99%', '52.56 minutes', '4.32 minutes'],
     ['99.999%', '5.26 minutes', '25.9 seconds']],
    col_widths=[4, 6, 6]
)

doc.add_heading('SDLC Phases', level=2)
add_table(
    ['Phase', 'Key Doc', 'Legal Risk', 'DPA Link'],
    [['Requirements', 'Specification', 'Incomplete requirements', 'S.20(2)(a) identify risks'],
     ['Design', 'Architecture', 'No security-by-design', 'S.20(2)(b) establish safeguards'],
     ['Dev', 'Source code', 'IP, open source compliance', '\u2014'],
     ['Testing', 'Test reports', 'Inadequate testing', 'S.20(2)(c) verify safeguards'],
     ['Deployment', 'Deployment log', 'Unauthorised deployment', 'S.20(1) org measures'],
     ['Maintenance', 'Patch log', 'Unpatched vulns', 'S.20(2)(d) update safeguards']],
    col_widths=[3, 4, 4, 5]
)

doc.add_heading('ETA Key Sections for APIs', level=2)
add_table(
    ['Section', 'What It Says', 'Practical Use'],
    [['S.29', 'Provider not liable if merely providing access', 'Exemption lost if contract exists (S.29(2)(a))'],
     ['S.30', 'No liability for links to infringing content', 'API linking to third-party content'],
     ['S.31', 'Notice-and-takedown procedure', 'How to notify API provider'],
     ['S.32', 'No duty to monitor', 'Provider not required to police data proactively'],
     ['S.33', 'Territorial jurisdiction', 'Applies to acts inside or outside Uganda'],
     ['S.34', 'Jurisdiction of courts', 'Magistrate Grade 1/Chief Magistrate']],
    col_widths=[3, 7, 6]
)

doc.add_heading('DPA Key Sections for Cloud/SDLC', level=2)
add_table(
    ['Section', 'What It Says', 'Practical Use'],
    [['S.20', 'Technical and organisational security measures', 'API/SDLC/cloud security all mandated'],
     ['S.21', 'Controller-processor contract required', 'Cloud DPA is statutory, not optional'],
     ['S.22', 'Operator confidentiality', 'Provider staff confidentiality obligations'],
     ['Reg. 31', 'Publish security practices', 'Document API/cloud security measures'],
     ['Reg. 32', 'Specific security measures', 'Access control, encryption, monitoring'],
     ['Reg. 33', 'Breach notification to PDPO', 'API breach = mandatory notification']],
    col_widths=[3, 7, 6]
)

add_divider()

# ============================================================
# PART 8: GLOSSARY
# ============================================================
doc.add_heading('PART 8: GLOSSARY FOR LAWYERS', level=1)

glossary = [
    ('Agile', 'Software development in short cycles (sprints) with continuous feedback and adaptation.'),
    ('API', 'Rules allowing one software application to communicate with another. Like a waiter between customer and kitchen.'),
    ('API Endpoint', 'The specific URL where an API can be accessed. Like a specific desk in a government office.'),
    ('API Key', 'A string identifying a client application to an API server. Like a building access card.'),
    ('Cloud Computing', 'Delivering computing services over the internet, paying only for what you use.'),
    ('Data Localisation', 'Requirement that data be stored on servers within a specific country\'s borders.'),
    ('DPA (agreement)', 'Contract between data controller and processor specifying data protection obligations. Required by DPA S.21.'),
    ('DevOps', 'Development and operations combined into one team with continuous deployment.'),
    ('DPIA', 'Data Protection Impact Assessment \u2014 identifies and mitigates privacy risks before processing.'),
    ('HTTP Verbs', 'Standard methods (GET, POST, PUT, PATCH, DELETE) that tell an API what action to perform.'),
    ('IaaS', 'Infrastructure as a Service \u2014 renting raw computing resources from a cloud provider.'),
    ('JSON', 'Lightweight format for transmitting data \u2014 the most common API payload format.'),
    ('LLMOps', 'Lifecycle management for Large Language Models, with governance layers for bias, explainability, and security.'),
    ('OAuth 2.0', 'Standard protocol for authorising API access using temporary tokens.'),
    ('OWASP API Top 10', 'Industry-standard list of the 10 most critical API security risks (2023 edition).'),
    ('PaaS', 'Platform as a Service \u2014 renting a complete platform from a cloud provider.'),
    ('Payload', 'The actual data carried by an API request or response, typically in JSON or XML.'),
    ('REST', 'Design principles for building web APIs using standard HTTP methods and stateless communication.'),
    ('SaaS', 'Software as a Service \u2014 using a vendor\'s application over the internet.'),
    ('SDLC', 'Software Development Lifecycle \u2014 Requirements, Design, Dev, Testing, Deployment, Maintenance.'),
    ('SLA', 'Service Level Agreement \u2014 contractual promise of service quality with measurable targets and remedies.'),
    ('Sovereign Cloud', 'Cloud infrastructure within a country\'s borders, under that country\'s laws.'),
    ('Token', 'Temporary digital pass to authenticate API requests. Like a visitor\'s badge.'),
    ('UGHub', 'Uganda\'s national API gateway operated by NITA-U, connecting 135+ government entities.'),
    ('UNCITRAL Notes', 'UN best-practice guidelines for drafting cloud computing contracts (2019).'),
    ('Waterfall', 'Traditional software development where each phase completes before the next begins.'),
]

for term, defn in glossary:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{term}: ')
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(defn)
    r2.font.size = Pt(10)

# Footer
add_divider()
add_para('End of Week 4 Reading Notes', italic=True, size=10, color='888888', align=WD_ALIGN_PARAGRAPH.CENTER)
add_para('Corrected statutory references: DPA Section 20 (not Section 22). DPA Section 19 (cross-border transfers; not Section 26). ETA Section 7 (originals; not Section 6). DPA Regulations 31\u201333 (not Regulation 22). UCC Act applies conditionally \u2014 only to licensed operators and data communication service providers.', size=9, color='999999', align=WD_ALIGN_PARAGRAPH.CENTER)

# Save
output = r'C:\Users\DELL\research\Week 4 - APIs, Cloud Computing & SDLC.docx'
doc.save(output)
print(f'Saved: {output}')

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

doc = Document()

# -- Global default font --
style = doc.styles['Normal']
font = style.font
font.name = 'Bookman Old Style'
font.size = Pt(12)

# Also set heading styles
for level in range(1, 5):
    hstyle = doc.styles[f'Heading {level}']
    hfont = hstyle.font
    hfont.name = 'Bookman Old Style'
    hfont.color.rgb = RGBColor(0, 0, 0)
    if level == 1:
        hfont.size = Pt(16)
        hfont.bold = True
    elif level == 2:
        hfont.size = Pt(14)
        hfont.bold = True
    elif level == 3:
        hfont.size = Pt(13)
        hfont.bold = True
    elif level == 4:
        hfont.size = Pt(12)
        hfont.bold = True
        hfont.italic = True

def add_para(text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Bookman Old Style'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def add_table_row(table, cells, bold=False):
    row = table.add_row()
    for i, txt in enumerate(cells):
        cell = row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(txt)
        run.font.name = 'Bookman Old Style'
        run.font.size = Pt(11)
        run.bold = bold
    return row

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('Module 1: Digital Technology Fundamentals', level=1)
subtitle = doc.add_heading('Week 2: Internet Architecture (TCP/IP, DNS, HTTP/HTTPS, Web Applications)', level=2)

doc.add_paragraph('─' * 60)

# ============================================================
# LEARNING OBJECTIVES
# ============================================================
doc.add_heading('Learning Objectives', level=3)
objectives = [
    'Map the TCP/IP 5-layer model to data flows in digital forensics and determine at which layer an interception, breach, or service failure occurred.',
    'Analyse the DNS resolution chain to establish jurisdiction over domain registration data and trace content distribution paths.',
    'Evaluate HTTP/HTTPS communications to distinguish between encrypted and unencrypted data for purposes of lawful interception under RICA and the DPA.',
    'Apply Uganda\'s content and consumer protection regulations to internet-based services, including OTT platforms, web applications, and cloud computing.'
]
for i, obj in enumerate(objectives, 1):
    add_para(f'{i}. {obj}')

doc.add_paragraph('─' * 60)

# ============================================================
# 2.1 THE INTERNET AS A NETWORK OF NETWORKS
# ============================================================
doc.add_heading('2.1 The Internet as a Network of Networks', level=2)

doc.add_heading('2.1.1 The 5-Layer TCP/IP Model', level=3)
add_para('The Internet does not use the 7-layer OSI model in practice. Instead, it operates on a 5-layer TCP/IP model:')

table = doc.add_table(rows=1, cols=4, style='Table Grid')
header = table.rows[0]
for i, txt in enumerate(['Layer', 'TCP/IP Model', 'Protocols', 'Role']):
    cell = header.cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name = 'Bookman Old Style'
    run.font.size = Pt(11)
    run.bold = True

rows_data = [
    ['5', 'Application', 'HTTP, HTTPS, DNS, SMTP, FTP', 'User-facing protocols; data originates here'],
    ['4', 'Transport', 'TCP, UDP', 'End-to-end delivery; port numbers, segmentation'],
    ['3', 'Network', 'IP (IPv4, IPv6)', 'Routing; logical addressing across networks'],
    ['2', 'Link', 'Ethernet, Wi-Fi (802.11)', 'Physical addressing (MAC); local frame delivery'],
    ['1', 'Physical', 'Copper, fibre, radio', 'Raw bit transmission'],
]
for r in rows_data:
    add_table_row(table, r)

doc.add_paragraph()
add_para('Encapsulation is the process by which each layer adds its own header (and sometimes trailer):')
add_para('Application Data -> TCP Header + App Data -> IP Header + TCP + App -> Ethernet Header + IP + TCP + App + Trailer', italic=True)

add_para('For the legal practitioner, each header reveals different jurisdictional information:')
add_para('\u2022 The IP header (Layer 3) reveals source/destination IP addresses, which can be geolocated.')
add_para('\u2022 The TCP header (Layer 4) reveals port numbers, identifying the type of application (port 80 = HTTP, port 443 = HTTPS, port 53 = DNS).')
add_para('\u2022 The application data (Layer 5) is the actual content \u2014 which may be encrypted (HTTPS) or plaintext (HTTP).')

doc.add_heading('2.1.2 The Network Edge vs. the Network Core', level=3)
add_para('Network Edge: End systems (hosts) where applications run \u2014 smartphones, laptops, servers, IoT devices.')
add_para('Network Core: The mesh of routers and links that move data between edge devices.')
add_para('Legal significance: Lawful interception under RICA targets specific points in the network. The edge/core distinction determines whether data is "in transit" (core) or "at rest on a device" (edge), affecting both the warrant requirement and admissibility.')

doc.add_heading('2.1.3 Packet Switching vs. Circuit Switching', level=3)
add_para('The Internet uses packet switching: data is broken into packets, each routed independently. This creates jurisdictional complexity for cross-border data claims under DPA \u00a719, because packets from Kampala to Nairobi may transit through London or Dubai.')

doc.add_paragraph('─' * 60)

# ============================================================
# 2.2 DNS
# ============================================================
doc.add_heading('2.2 The Domain Name System (DNS)', level=2)

doc.add_heading('2.2.1 The Problem DNS Solves', level=3)
add_para('DNS is the Internet\u2019s directory service that translates human-friendly hostnames (www.google.com) into machine-readable IP addresses (142.250.190.4). It is (1) a distributed database implemented in a hierarchy of DNS servers, and (2) an application-layer protocol running over UDP port 53.')

doc.add_heading('2.2.2 The DNS Hierarchy', level=3)
add_para('No single DNS server holds all mappings. The hierarchy is:')
add_para('Root DNS Servers \u2192 TLD Servers (.com, .ug, .ke) \u2192 Authoritative DNS Servers \u2192 Local DNS Server (ISP-provided)')

add_para('Root servers provide IP addresses of TLD servers. TLD servers handle top-level domains. Authoritative servers hold actual DNS records. Local DNS servers act as proxies, forwarding queries from user hosts into the hierarchy.')

doc.add_heading('2.2.3 DNS Resolution Process', level=3)
add_para('When a user in Kampala types www.example.com into a browser:')
add_para('1. Browser calls gethostbyname() (client side of DNS).')
add_para('2. Query sent to local DNS server (e.g., MTN Uganda).')
add_para('3. Local DNS queries a root server \u2192 gets TLD server for .com.')
add_para('4. Local DNS queries .com TLD server \u2192 gets authoritative server for example.com.')
add_para('5. Local DNS queries authoritative server \u2192 gets IP address.')
add_para('6. Browser initiates TCP connection to that IP address.')
add_para('Caching: DNS responses are cached at the local DNS server based on Time-To-Live (TTL) values.')

doc.add_heading('2.2.4 DNS Services Beyond Address Translation', level=3)
add_para('\u2022 Host aliasing: A complicated canonical hostname can have simpler alias names (relevant to phishing/fraud cases).')
add_para('\u2022 Mail server aliasing: MX records allow mail/web servers to share hostnames (relevant to tracing email origin).')
add_para('\u2022 Load distribution: DNS rotates IP addresses among replicated servers (used by Akamai, Cloudflare).')

doc.add_heading('2.2.5 The .ug ccTLD and Ugandan Domain Administration', level=3)
add_para('The .ug country-code top-level domain is administered by the Uganda Internet Exchange Point (UIXP) under the authority of UCC. WHOIS lookups reveal registrant information relevant to content disputes. UCC Content Regulations, 2019, Regulation 7 requires operators to retain content records for at least 60 days \u2014 applying to internet-based content services as well (Regulation 5).')

doc.add_paragraph('─' * 60)

# ============================================================
# 2.3 HTTP AND HTTPS
# ============================================================
doc.add_heading('2.3 HTTP and HTTPS', level=2)

doc.add_heading('2.3.1 The Web\u2019s Application-Layer Protocol', level=3)
add_para('HTTP is a stateless, text-based request-response protocol running over TCP (port 80 for HTTP, port 443 for HTTPS).')

doc.add_heading('2.3.2 HTTP Request-Response Cycle', level=3)
add_para('Client (Browser) sends request: GET /index.html HTTP/1.1 with Host, User-Agent, Cookie headers.')
add_para('Server responds: HTTP/1.1 200 OK with Content-Type, Set-Cookie, body.')

table2 = doc.add_table(rows=1, cols=3, style='Table Grid')
for i, txt in enumerate(['HTTP Method', 'Purpose', 'Legal Relevance']):
    cell = table2.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name = 'Bookman Old Style'
    run.font.size = Pt(11)
    run.bold = True
for r in [
    ['GET', 'Retrieve a resource', 'Access logs show what content was requested and when'],
    ['POST', 'Submit data to be processed', 'Form submissions, login credentials, payment data'],
    ['PUT', 'Upload/replace a resource', 'Content uploads to servers'],
    ['DELETE', 'Remove a resource', 'Evidence of intentional data destruction'],
    ['PATCH', 'Partial modification', 'Data alteration records'],
]:
    add_table_row(table2, r)

doc.add_heading('2.3.3 HTTP Status Codes', level=3)
table3 = doc.add_table(rows=1, cols=4, style='Table Grid')
for i, txt in enumerate(['Code', 'Class', 'Meaning', 'Example']):
    cell = table3.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name = 'Bookman Old Style'
    run.font.size = Pt(11)
    run.bold = True
for r in [
    ['1xx', 'Informational', 'Request received, processing', '101 Switching Protocols'],
    ['2xx', 'Success', 'Request understood', '200 OK, 201 Created'],
    ['3xx', 'Redirection', 'Further action needed', '301 Moved Permanently'],
    ['4xx', 'Client Error', 'Request cannot be fulfilled', '400 Bad, 401 Unauthorized, 403 Forbidden, 404 Not Found'],
    ['5xx', 'Server Error', 'Server failed', '500 Internal Server Error, 502 Bad Gateway'],
]:
    add_table_row(table3, r)

doc.add_heading('2.3.4 Cookies and State Management', level=3)
add_para('HTTP is stateless. Cookies create stateful sessions via Set-Cookie (server to browser) and Cookie (browser to server). Legal significance: session hijacking via intercepted cookies is directly relevant to Computer Misuse Act \u00a712 (unauthorised access) and \u00a714 (access with intent).')

doc.add_heading('2.3.5 HTTPS and TLS', level=3)
add_para('HTTPS encrypts the entire HTTP message using TLS, providing encryption, authentication (digital certificates), and integrity.')
add_para('Critical note for lawful interception: Under RICA, interception requires a warrant (RICA \u00a75). However, HTTPS means that even with a warrant, the content of the communication may remain encrypted and inaccessible unless the interceptor also obtains TLS session keys.')

doc.add_paragraph('─' * 60)

# ============================================================
# 2.4 WEB APPLICATIONS AND CLOUD SERVICES
# ============================================================
doc.add_heading('2.4 Web Applications and Cloud Services', level=2)

doc.add_heading('2.4.1 From Static Pages to Web Applications', level=3)
add_para('Modern web applications involve client-side processing (JavaScript), server-side processing (application servers, databases), and APIs (RESTful, GraphQL).')

doc.add_heading('2.4.2 Cloud Computing Models', level=3)
table4 = doc.add_table(rows=1, cols=4, style='Table Grid')
for i, txt in enumerate(['Model', 'Description', 'Example', 'Legal Implications']):
    cell = table4.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name = 'Bookman Old Style'
    run.font.size = Pt(11)
    run.bold = True
for r in [
    ['IaaS', 'Virtualised computing resources', 'AWS EC2, Google Compute Engine', 'Customer controls OS; provider controls physical infrastructure'],
    ['PaaS', 'Platform for deploying applications', 'Google App Engine, Heroku', 'Provider manages OS/runtime; customer manages app code'],
    ['SaaS', 'Ready-to-use software over web', 'Google Workspace, Microsoft 365', 'Provider controls everything; customer only uses the app'],
]:
    add_table_row(table4, r)

doc.add_heading('2.4.3 Data Sovereignty and Cloud Jurisdiction', level=3)
add_para('DPA \u00a719 requires that where personal data is processed or stored outside Uganda, the recipient country must have adequate measures at least equivalent to the DPA, or the data subject must consent. A Ugandan company using AWS must verify adequacy of the host country or obtain explicit consent.')

doc.add_paragraph('─' * 60)

# ============================================================
# 2.5 UGANDA'S REGULATORY FRAMEWORK
# ============================================================
doc.add_heading("2.5 Uganda's Regulatory Framework for the Internet", level=2)

doc.add_heading('2.5.1 The Uganda Communications Act (UCA) and UCC', level=3)
add_para('UCC is the primary regulator of communications services under the Uganda Communications Act, 2013. Powers include ISP licensing, content regulation (Content Regulations 2019), consumer protection (Consumer Protection Regulations 2019), quality of service monitoring, and equipment type approval.')

doc.add_heading('2.5.2 UCC Content Regulations, 2019', level=3)
add_para('Apply to "all content in telecommunications, data and radio communications and broadcasting and postal communications" (Reg. 1(2)). Key provisions:')
add_para('\u2022 Reg. 5: Defines content services broadly, covering internet-based content distribution.')
add_para('\u2022 Reg. 7: Operators must retain records for at least 60 days; records must be "complete, authentic and original."')
add_para('\u2022 Reg. 8: Prohibits offensive language, explicit sexual content, incitement, and material contrary to public morality.')
add_para('\u2022 Reg. 38: Prohibits broadcasting material relating to private affairs without compelling public interest.')
add_para('\u2022 Reg. 45: Offence punishable by fine not exceeding 48 currency points (UGX 960,000) or imprisonment up to 2 years.')

doc.add_heading('2.5.3 UCC Consumer Protection Regulations, 2019', level=3)
add_para('Key provisions:')
add_para('\u2022 Reg. 6: Rights of consumers \u2014 access, choice, accurate billing, redress.')
add_para('\u2022 Reg. 10: Prohibited advertising \u2014 false, misleading, bait-and-switch.')
add_para('\u2022 Reg. 13: Prohibits denial of access except for non-payment or just cause.')
add_para('\u2022 Reg. 16: Protection of consumer information \u2014 must be lawfully collected, processed for identified purposes, accurate, protected against improper disclosure.')
add_para('\u2022 Reg. 18: Operators must protect consumers from spam, scams, unsolicited calls, harmful content; opt-out must be free.')
add_para('\u2022 Reg. 24-25: Service Level Agreements must be submitted to UCC; bundled services remain operator\u2019s responsibility.')

doc.add_heading('2.5.4 The Regulation of Interception of Communications Act (RICA)', level=3)
add_para('RICA governs lawful interception:')
add_para('\u2022 Section 2: "Interception" means listening to, recording, monitoring, or acquiring content without knowledge of communicants.')
add_para('\u2022 Section 5: No interception without a High Court warrant specifying target, duration, and scope.')
add_para('\u2022 Technical tension: HTTPS encryption means even with a warrant, content may remain inaccessible.')

doc.add_heading('2.5.5 Cross-Border Data Transfers Under the DPA', level=3)
add_para('DPA \u00a719 requires adequacy in the recipient country or data subject consent. The PDPO has not published a list of "adequate" jurisdictions, creating legal uncertainty.')
add_para('Landmark case: Ssekamwa Frank & 3 Ors v. Google LLC (PDPO, July 2025) \u2014 Google ordered to register with PDPO; the PDPO affirmed extraterritorial application of the DPA and required documented legal basis for cross-border transfers.')

doc.add_heading('2.5.6 The Proposed Single Digital Media Law (2026)', level=3)
add_para('In April 2026, the Government announced a draft law consolidating the Communications Act, Press and Journalists Act, Computer Misuse Act, and DPA. Key provisions: unified licensing for digital platforms; local representation for offshore platforms; mandatory content moderation/takedown (24-48 hrs); metadata retention for ISPs; enhanced cross-border data transfer controls; AI/automated decisioning rules. The bill had not been gazetted as of mid-2026.')

doc.add_heading('2.5.7 Regional and Continental Frameworks', level=3)
add_para('EAC Framework for Cyberlaws (2009/2010): First regional cyberlaw framework in Africa. Uganda\u2019s ETA, CMA, and Electronic Signatures Act (all 2011) were enacted in direct response. A Cross-Border Data Flows Framework was validated in June 2026 under EARDIP.')
add_para('Malabo Convention (2014, in force 2023): AU\u2019s binding treaty on data protection, electronic transactions, and cybersecurity. Uganda has NOT signed or ratified. Only Rwanda in East Africa has ratified. Uganda\u2019s DPA already meets its core standards.')

doc.add_paragraph('─' * 60)

# ============================================================
# 2.6 LEGAL CHALLENGES
# ============================================================
doc.add_heading('2.6 Legal Challenges in Internet Regulation', level=2)

doc.add_heading('2.6.1 OTT Regulation', level=3)
add_para('Key issues: (1) Content liability \u2014 is an OTT provider a "content provider" under UCC Content Regulations? (2) Consumer protection \u2014 do Consumer Protection Regulations apply to OTT subscriptions? (3) Data localisation \u2014 does the DPA require OTT providers to store Ugandan user data locally?')

doc.add_heading('2.6.2 Content Moderation', level=3)
add_para('Tension between: (1) operator obligation to remove prohibited content (Content Regs Reg. 8); (2) user right to freedom of expression (Constitution Art. 29); (3) absence of statutory safe harbour for user-generated content (unlike US Section 230 or EU Digital Services Act).')

doc.add_heading('2.6.3 Data Localisation and the Adequacy Gap', level=3)
add_para('Challenges: (1) Uganda has limited local cloud infrastructure; (2) local hosting is more expensive; (3) globally distributed services lack local redundancy.')

doc.add_paragraph('─' * 60)

# ============================================================
# 2.7 WEEKLY PRACTICE TASK
# ============================================================
doc.add_heading('2.7 Weekly Practice Task: Technical Deposition \u2014 Internet Layer Tracing', level=2)

add_para('The Scenario:')
add_para('Your client, a Ugandan e-commerce company, suffered a data breach. Customer payment data was intercepted during transmission. The logs show: (1) browser connected to https://www.shopuganda.com (port 443); (2) DNS resolution via MTN Uganda\u2019s local DNS server; (3) server hosted on AWS EC2 (Cape Town); (4) traffic routed through an unknown IP in Dubai before reaching the Cape Town server.')

add_para('Your Task:')
add_para('Draft a Technical Deposition Questionnaire (max 10 questions) directed at the IT security officer to isolate: (1) which TCP/IP layer the interception occurred at; (2) whether HTTPS/TLS was properly implemented; (3) whether DNS was compromised; (4) regulatory implications under DPA, UCC Consumer Protection Regs, and RICA.')
add_para('Sample question: "Can you confirm whether the TLS certificate presented by the server to the customer\u2019s browser was valid, self-signed, or mismatched, and whether the browser generated a certificate warning?"')

doc.add_paragraph('─' * 60)

# ============================================================
# CHAPTER SUMMARY
# ============================================================
doc.add_heading('Chapter Summary', level=2)

table5 = doc.add_table(rows=1, cols=2, style='Table Grid')
for i, txt in enumerate(['Concept', 'Key Takeaway for Legal Practice']):
    cell = table5.rows[0].cells[i]
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(txt)
    run.font.name = 'Bookman Old Style'
    run.font.size = Pt(11)
    run.bold = True
for r in [
    ['5-Layer TCP/IP Model', 'Data interception occurs at a specific layer; each layer reveals different jurisdictional information'],
    ['DNS Resolution', 'Chain of queries establishes jurisdiction over domain registration and identifies spoofing'],
    ['HTTP vs. HTTPS', 'Encrypted vs. plaintext determines what data is accessible under a RICA warrant'],
    ['Web Apps and Cloud', 'Cloud jurisdiction depends on server location and DPA \u00a719 adequacy'],
    ['UCC Content Regulations', 'Apply broadly to internet content; 60-day record retention'],
    ['UCC Consumer Protection', 'ISPs and OTT operators must protect consumer info, provide SLAs'],
    ['DPA \u00a719', 'Cross-border transfers require adequacy or consent'],
    ['Single Digital Media Law 2026', 'Proposed consolidated framework: licensing, moderation, metadata retention'],
    ['EAC Cyberlaw Framework', 'Uganda\'s ETA/CMA/DPA derive from EAC harmonisation'],
    ['Malabo Convention', 'Continental benchmark; Uganda not ratified but DPA meets core standards'],
]:
    add_table_row(table5, r)

doc.add_paragraph('─' * 60)

# ============================================================
# REFERENCES
# ============================================================
doc.add_heading('References', level=2)
refs = [
    'Kurose and Ross, Computer Networking: A Top-Down Approach (8th ed., Pearson, 2021), Sections 1.5, 2.2, 2.4.',
    'The Uganda Communications (Content) Regulations, 2019 (S.I. 2019 No. 91).',
    'The Uganda Communications (Consumer Protection) Regulations, 2019 (S.I. 2019 No. 87).',
    'The Data Protection and Privacy Act, No. 9 of 2019 (Uganda), Section 19.',
    'The Regulation of Interception of Communications Act, No. 19 of 2010 (Uganda), Sections 2 and 5.',
    'Ssekamwa Frank & 3 Ors v. Google LLC, PDPO Complaint No. 08/11/24/6683 (18 July 2025).',
    'Isaac Ssejjombwe, "Govt drafts single law to govern media, digital space," Daily Monitor, 27 April 2026.',
    'UNCTAD, "Harmonizing Cyberlaws and Regulations: The Experience of the East African Community" (2012).',
    'African Union Convention on Cyber Security and Personal Data Protection (Malabo Convention, 2014, in force 2023).',
    'Chambers and Partners, TMT 2026 \u2014 Uganda (Global Practice Guides, February 2026).',
]
for ref in refs:
    add_para(ref, size=11)

doc.add_paragraph('─' * 60)

# ============================================================
# FURTHER READING
# ============================================================
doc.add_heading('Further Reading', level=2)
further = [
    'Kurose and Ross, Computer Networking: A Top-Down Approach (8th ed.) \u2014 Sections 1.5, 2.2, 2.4.',
    'UCC Content Regulations, 2019 (S.I. 2019 No. 91).',
    'UCC Consumer Protection Regulations, 2019 (S.I. 2019 No. 87).',
    'Data Protection and Privacy Act, No. 9 of 2019 (Uganda), Section 19.',
    'Regulation of Interception of Communications Act, No. 19 of 2010 (Uganda).',
    'RFC 1034/1035 (DNS), RFC 7230-7235 (HTTP/1.1), RFC 8446 (TLS 1.3).',
    'UNCTAD, "Harmonizing Cyberlaws and Regulations: The Experience of the EAC" (2012).',
    'Malabo Convention (AU, 2014) \u2014 Articles 8-22 on data protection.',
    'DLA Piper Africa, "Uganda Data Protection Regulator Clarifies Compliance Requirements" (July 2025).',
]
for f_item in further:
    add_para(f_item, size=11)

# ============================================================
# SAVE
# ============================================================
doc.save('C:\\Users\\DELL\\research\\Module 1 - Digital Technology Fundamentals - Week 2 - Internet Architecture.docx')
print('Saved successfully')

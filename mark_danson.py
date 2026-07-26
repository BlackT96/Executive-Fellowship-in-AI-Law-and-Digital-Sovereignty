import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_comment(doc, text, color_rgb=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.bold = True
    if color_rgb:
        run.font.color.rgb = color_rgb
    return p

def add_normal(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_mark(doc, label, marks_awarded, marks_total, comment=""):
    p = doc.add_paragraph()
    run = p.add_run(f"[{label}: {marks_awarded}/{marks_total}]")
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 100, 0)
    if comment:
        run2 = p.add_run(f" {comment}")
        run2.font.size = Pt(10)
        run2.italic = True
    return p

def add_correction(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(f"CORRECTION: {text}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(180, 0, 0)
    run.italic = True
    return p

doc = docx.Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
pf = style.paragraph_format
pf.space_after = Pt(4)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EXECUTIVE FELLOWSHIP IN AI LAW & DIGITAL SOVEREIGNTY")
run.bold = True
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Module 1 - Week 3 - Databases")
run.bold = True
run.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CANDIDATE: DANSON TWESIGOMWE")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 0, 120)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Formative Assessment - 50 Marks")
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("FINAL SCORE: 33.5 / 50 (67%) - GRADE C+")
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0, 100, 0)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Marking Guide Key: [Marks Awarded/Marks Available] | CORRECTION in red")
run.font.size = Pt(10)
run.italic = True

doc.add_paragraph()

# ===== QUESTION 1 =====
add_comment(doc, "QUESTION 1: Which Record to Believe? (12 marks)")
doc.add_paragraph()

add_mark(doc, "Q1(a) - ACID/BASE Explanation", 3, 4, "Good plain-English explanation. Minor imprecision in phrasing.")
add_normal(doc, "Danson's Answer (Q1a): 'PostgreSQL is a database management system that applies to relational database while Cassandra operates non-relational databases... ACID rules... all-or-nothing strategy... Cassandra distributes records across computers to gather the final record which can take time.'")
add_correction(doc, "The phrase 'Cassandra only captures the record of the transaction and not the transaction itself' is technically imprecise. Cassandra captures transaction data but does not guarantee when all nodes will agree on the state due to eventual consistency. Better: 'Cassandra stores the same data but allows temporary differences between nodes until synchronisation completes.'")
doc.add_paragraph()

add_mark(doc, "Q1(b) - Reliability Analysis", 2.5, 4, "Correct choice (PostgreSQL). Good ETA references (s.7(2), s.8(3), s.8(5)). Weakness: overstates PostgreSQL as 'conclusive scientific evidence' - no record is conclusive. Also 'Cassandra does not operate a single node' is inaccurate.")
add_normal(doc, "Danson's Answer (Q1b): 'PostgreSQL record is the most reliable because it relies on ACID... once the transaction is found to have been recorded, it is conclusive scientific evidence...'")
add_correction(doc, "Avoid 'conclusive scientific evidence' - electronic records carry presumptions under ETA s.8(5) but these are rebuttable. The correct framing: PostgreSQL's ACID guarantees make it more probative, not conclusive. Also: 'Cassandra does not operate a single node' is misleading - it operates on multiple individual nodes. Better: 'Cassandra is distributed across nodes; a query against one node may not reflect the full system state.'")
doc.add_paragraph()

add_mark(doc, "Q1(c) - BigQuery Analysis", 2.5, 4, "Correctly identifies BigQuery is serverless (affecting s.8(5) presumption). Notes query by own staff, not in ordinary course. Missing: Dremel/Colossus audit trail discussion, ETL reconciliation rules.")
add_normal(doc, "Danson's Answer (Q1c): 'BigQuery is serverless so we shall not have proof of the proper working of the computer... query was by our own staff... data cannot be taken to have been stored in the ordinary course.'")
add_correction(doc, "Good point on serverless architecture. However, also discuss: (1) BigQuery's Dremel query engine preserves execution logs; (2) Colossus provides immutable storage; (3) ETL reconciliation rules determine how conflicts were resolved. These affect weight, not admissibility.")
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Q1 Total: 8 / 12")
run.bold = True
doc.add_paragraph()

# ===== QUESTION 2 =====
add_comment(doc, "QUESTION 2: Admissibility Challenge (14 marks)")
doc.add_paragraph()

add_mark(doc, "Q2(a) - Objection 1 (Originality)", 4, 5, "Strong answer. Correctly notes s.6 is not relevant to originality - shows independent thinking. Good use of s.7(1)(b), s.8(1)(a), s.8(1)(c), s.8(3).")
add_normal(doc, "Danson's Answer (Q2a): 'Note: I do not find the relevance of section 6... Section 7 of the ETA which provides for authenticity of a data message... Section 8(1)(c) precludes this honourable court from rejecting admissibility because it is not in original form.'")
add_correction(doc, "You are correct that s.6 (electronic signatures) is not directly relevant to originality. Better approach: note s.6 deals with signatures, but the court should look to s.7 (authenticity of data message) for originality. Your instinct is right; presentation could be more diplomatic. Excellent use of s.8(1)(c) - this is the key provision defeating the best evidence objection.")
doc.add_paragraph()

add_mark(doc, "Q2(b) - Objection 2 (Hearsay)", 4, 5, "Excellent. Cites R v Spiby (1990) - correct authority. Correctly distinguishes machine-generated from human statements.")
add_normal(doc, "Danson's Answer (Q2b): 'R v Spiby (1990)... information recorded by a computer or machine without human intervention falls outside the scope of hearsay rules... the analyst only ran a query which provided what was already recorded.'")
add_correction(doc, "Strong. One refinement: you could strengthen by noting that even under Evidence Act s.59 (oral evidence must be direct), machine-generated records are real evidence, not testimonial statements. The query act is equivalent to opening a file cabinet.")
doc.add_paragraph()

add_mark(doc, "Q2(c) - Objection 3 (ETL Pipeline)", 3, 4, "Good technical understanding. Explains ETL, Dremel, Colossus, immutability, audit logs. Could be more specific on audit log contents.")
add_normal(doc, "Danson's Answer (Q2c): 'Dremel receives data and reviews it in seconds by dividing it among many devices... Colossus stores copies of all reviewed data... data is immutable... BigQuery audit logs can reveal who entered what, who queried what.'")
add_correction(doc, "Add: BigQuery's INFORMATION_SCHEMA provides query execution metadata; Cloud Audit Logs track admin/data access; time-travel allows querying historical data. These are specific discoverable records.")
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Q2 Total: 11 / 14")
run.bold = True
doc.add_paragraph()

# ===== QUESTION 3 =====
add_comment(doc, "QUESTION 3: Cross Examination (12 marks)")
doc.add_paragraph()

add_mark(doc, "Q3(a) - Strategy (5-Step Framework)", 1.5, 3, "Reasonable content but does NOT explicitly follow the 5-step problem-solving framework as instructed.")
add_normal(doc, "Danson's Answer (Q3a): 'Identifying the database type and DBMS... Consistency model... Statutory compliance... Challenge.'")
add_correction(doc, "The question required you to 'use the 5-step problem-solving framework from your reading notes (Part 5).' Your answer should explicitly label: Step 1 (Issue), Step 2 (Law), Step 3 (Application), Step 4 (Strengths/Weaknesses), Step 5 (Strategy). See Ingrid's answer for the correct format.")
doc.add_paragraph()

add_mark(doc, "Q3(b) - Four Cross-Examination Questions", 3.5, 4, "Four well-targeted questions with clear purposes. Q3 (number of servers) could be sharper.")
add_normal(doc, "Danson's Questions: Q1 'On which database do you record agents payments?' Q2 'What is the databases consistency model?' Q3 'How many servers does this database operate on?' Q4 'How often do you conduct active repair on your Cassandra system?'")
add_correction(doc, "Q3 could be sharper: instead of 'how many servers,' ask 'Was the query run against a primary node or a read replica?' This directly targets the replication lag. Q4 on active repair is excellent - it targets the s.8(5)(a) presumption.")
doc.add_paragraph()

add_mark(doc, "Q3(c) - Amongin Steps", 2, 3, "Identifies Steps 4, 5, 6, 10. Applies them to facts. Step 10 application is confused.")
add_normal(doc, "Danson's Amongin: 'Step 4. Method of storing data... Step 5. Reliability of computer programs... Step 6. Measures to verify accuracy... Step 10. Independent third party should achieve same results.'")
add_correction(doc, "Your Step 10 application states: 'When this record is subjected to acid properties, it cannot stand the consistencies' - Step 10 asks whether an independent third party running the same query would get the same result. The correct argument: if an expert queried all three Cassandra nodes simultaneously, results would differ due to replication lag, proving the printout is unreliable.")
doc.add_paragraph()

add_mark(doc, "Q3(d) - Answer to Judge on Weight", 2, 2, "Excellent. Cites s.8(4) correctly. Clear, concise, persuasive.")
add_normal(doc, "Danson's Answer: 'Section 8(4) provides that the court shall regard the reliability of the manner in which the data message was generated... Cassandra records implement BASE properties which do not guarantee consistency.'")
add_correction(doc, "Perfect. No corrections needed.")
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Q3 Total: 9 / 12")
run.bold = True
doc.add_paragraph()

# ===== QUESTION 4 =====
add_comment(doc, "QUESTION 4: Data Retention and Deletion (12 marks)")
doc.add_paragraph()

add_mark(doc, "Q4(a) - Deletion Deprivation", 1.5, 4, "Weakest answer. Identifies what each database stores but then contradicts himself by saying PostgreSQL captured everything.")
add_normal(doc, "Danson's Answer (Q4a): 'PostgreSQL keeps all data for all transaction records... Cassandra keeps customer registration data, agent commission records, and transaction history queries. Theres no record that PostgreSQL didnt capture.'")
add_correction(doc, "You correctly IDENTIFIED that Cassandra stores different data (registration, commissions, transaction history) but incorrectly CONCLUDED that PostgreSQL captured everything. The fact pattern states the Mbarara node recorded a UGX 500,000 commission reversal at 14:28 that never propagated. That reversal existed ONLY in Cassandra and was purged after 60 days. This is exactly the evidence the agent lost.")
doc.add_paragraph()

add_mark(doc, "Q4(b) - 60-Day Retention Compliance", 2, 4, "References NPS Act (good) and s.9 ETA (good). Missing: Limitation Act discussion (question specifically asked for it).")
add_normal(doc, "Danson's Answer (Q4b): 'Under the National Payment Systems Act cap 59, a service provider must keep records for at least 10 years. Section 9(1)(a) ETA provides records must be accessible.'")
add_correction(doc, "Two issues: (1) The National Payment Systems Act is Act No. 10 of 2020, not 'Cap. 59'. (2) The question asks you to 'consider the Limitation Act Cap. 80 period for contractual claims (6 years).' You did not address this. The argument: if a contractual claim can be brought within 6 years (Limitation Act), a 60-day retention policy is grossly inadequate under s.9 ETA.")
doc.add_paragraph()

add_mark(doc, "Q4(c) - Retention Policy Draft", 2, 4, "Correctly cites DPA s.18(1)(a) (not s.14 - good catch). Answer too short, uniform recommendation.")
add_normal(doc, "Danson's Answer (Q4c): 'Section 18(1)(a) DPA prohibits retaining data longer than necessary... Section 63 NPS Act requires 10 years. I recommend 10 years for all systems.'")
add_correction(doc, "Good statutory identification (s.18(1)(a) DPA, NPS Act). However: (1) A single 10-year policy ignores data minimization under DPA s.18 - Cassandra stores registration data that may not need 10-year retention. (2) The question asks you to explain why different retention periods may be justified. Different systems store different data types with different legal requirements.")
doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Q4 Total: 5.5 / 12")
run.bold = True
doc.add_paragraph()

# ===== OVERALL =====
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("=" * 50)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ASSESSMENT SUMMARY")
run.bold = True
run.font.size = Pt(13)

summary_items = [
    ("Question 1: Which Record to Believe?", "8 / 12", "66%"),
    ("Question 2: Admissibility Challenge", "11 / 14", "79%"),
    ("Question 3: Cross Examination", "9 / 12", "75%"),
    ("Question 4: Data Retention", "5.5 / 12", "46%"),
    ("TOTAL", "33.5 / 50", "67% - C+"),
]

for item, score, pct in summary_items:
    p = doc.add_paragraph()
    run = p.add_run(f"{item}: ")
    run.bold = True
    run = p.add_run(f"{score} ({pct})")

doc.add_paragraph()
add_comment(doc, "KEY STRENGTHS:")
for s in ["Good grasp of ACID vs BASE concepts", "Correct use of R v Spiby for hearsay distinction", "Strong cross-examination questions with clear purposes", "Noticed s.6 (signatures) is irrelevant to originality - independent thinking"]:
    add_normal(doc, f"  - {s}")

doc.add_paragraph()
add_comment(doc, "KEY WEAKNESSES:")
for s in ["Q4: Failed to connect Cassandras distinct data to the deletion argument", "Q4: Ignored Limitation Act question entirely", "Q3(a): Did not follow the 5-step framework structure as instructed", "Overstated language ('conclusive scientific evidence') undermines credibility"]:
    add_normal(doc, f"  - {s}")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("AI GENERATION ASSESSMENT: No significant AI generation detected.")
run.font.size = Pt(10)
run.italic = True

path = r'C:\Users\DELL\research\Week_3_Danson_Marked.docx'
doc.save(path)
print(f"Danson marked assessment saved to {path}")
